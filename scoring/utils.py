# app/utils.py  (ADD or REPLACE these sections)

from __future__ import annotations
import re
from typing import List, Dict, Any

# ---- Section weights (used by dynamic analyzer; your report table uses SECTION_MAX in views) ----
TECHNICAL_WEIGHTS = {
    "GitHub Profile": 27,
    "LinkedIn": 18,
    "Portfolio Website": 23,
    "Resume (ATS Score)": 23,
    "Certifications & Branding": 9,
    "LeetCode/DSA Skills": 9,
}
TECHNICAL_WEIGHTS = {
    "GitHub Profile": 27,
    "LinkedIn": 18,
    "Portfolio Website": 23,
    "Resume (ATS Score)": 23,
    "Certifications & Branding": 9,
    "LeetCode/DSA Skills": 9,
}

def get_grade_tag(score: float) -> str:
    """
    Generic label mapper used across the app.
    Accepts any 0..100-ish number and maps to a qualitative tag.
    """
    try:
        s = float(score)
    except Exception:
        s = 0.0
    if s >= 85: return "Excellent"
    if s >= 70: return "Good"
    if s >= 50: return "Average"
    return "Poor"

def calculate_dynamic_ats_score_v2(resume_text: str, github_username: str, leetcode_username: str, extracted_links: List[Dict[str, Any]]):
    """
    Deterministic, rubric-aligned dynamic scoring (technical).
    Inputs:
      - resume_text: str
      - github_username: str | ""
      - leetcode_username: str | ""
      - extracted_links: List[Dict] with keys at least {"url","domain","type","title","description"}
        NOTE: 'type' should be from your link extractor: "github","linkedin","leetcode","blog","video","notion","docs","other"
    Output:
      {
        "sections": {
           "<section>": {"score": int, "grade": str, "weight": TECHNICAL_WEIGHTS[section], "sub_criteria":[...]}
        },
        "overall_score_average": int,
        "overall_grade": str,
        "suggestions": [str, ...]
      }
    """
    def has_word(text, *words):
        t = (text or "").lower()
        return any(w.lower() in t for w in words)

    def distinct_links_of(types):
        # normalize set logic while preserving dicts
        bucket: Dict[str, Dict[str, Any]] = {}
        for l in extracted_links or []:
            t = (l.get("type") or "").lower()
            if t in {x.lower() for x in (types if isinstance(types, (set, list, tuple)) else [types])}:
                u = (l.get("url") or "").strip()
                if u and u not in bucket:
                    bucket[u] = l
        return list(bucket.values())

    def all_links_domain_contains(*needles):
        urls = [(l.get("url") or "") for l in (extracted_links or [])]
        needles_low = [n.lower() for n in needles]
        return any(any(n in u.lower() for n in needles_low) for u in urls)

    def grade_tag(score, max_points):
        try:
            return get_grade_tag(score)
        except Exception:
            pct = 0 if max_points <= 0 else (score / max_points) * 100
            if pct >= 85: return "Excellent"
            if pct >= 70: return "Good"
            if pct >= 50: return "Average"
            return "Poor"

    text_lower = (resume_text or "").lower()

    gh_links     = distinct_links_of({"github"})
    li_links     = distinct_links_of({"linkedin"})
    lc_links     = distinct_links_of({"leetcode"})
    blog_links   = distinct_links_of({"blog"})
    notion_links = distinct_links_of({"notion"})
    other_links  = distinct_links_of({"other"})
    demoish      = [l for l in (extracted_links or []) if has_word(l.get("url",""), "demo", "app.", "vercel", "netlify", "onrender", "cloudfront", "pages.dev", "github.io")]

    github_presence   = bool(github_username) or bool(gh_links)
    leetcode_presence = bool(leetcode_username) or bool(lc_links)
    linkedin_presence = bool(li_links)
    
    # Enhanced portfolio detection
    portfolio_links = other_links + blog_links + notion_links
    portfolio_presence = bool(portfolio_links) or (
        has_word(resume_text, "portfolio", "personal website", "project", "demo") and len(extracted_links or []) > 0
    )
    
    cert_presence = bool(re.search(r'\b(certification|certified|certificate|course)\b', resume_text or "", re.I))

    # ---------------- GitHub (0–27) ----------------
    gh_sub = []
    gh_score = 0
    if github_presence:
        gh_sub.append({"name":"Public link present","score":3,"weight":3,"insight":"GitHub link detected."})
        repo_link_count = sum(1 for l in gh_links if re.search(r"github\.com/[^/]+/[^/#?]+", (l.get("url") or ""), re.I))
        mentions_docs = has_word(resume_text, "readme", "docs", "documentation")
        mentions_tests = has_word(resume_text, "pytest", "unittest", "jest", "cypress", "tests", "ci", "github actions")
        multi_stack = sum(1 for k in ["python","javascript","typescript","java","go","rust","scala","kotlin","c++","c#","sql","docker","kubernetes","terraform","spark","airflow"] if k in text_lower) >= 5
        has_demo = len(demoish) > 0 or all_links_domain_contains("github.io")

        if repo_link_count == 0:
            base = 3
        elif repo_link_count == 1:
            base = 8
        elif repo_link_count == 2:
            base = 12
        else:
            base = 18

        if mentions_docs: base += 2
        if mentions_tests: base += 2
        if has_demo: base += 2
        if multi_stack: base += 3
        gh_score = max(2, min(27, base))

        gh_sub.extend([
            {"name":"Repo links detected","score":min(6, repo_link_count*2),"weight":6,"insight":f"{repo_link_count} repo link(s) found."},
            {"name":"Docs/README/tests/CI","score":(2 if mentions_docs else 0)+(2 if mentions_tests else 0),"weight":4,"insight":"Signals of quality and maintainability."},
            {"name":"Hosted demos","score":2 if has_demo else 0,"weight":2,"insight":"Live demo or GitHub Pages present."},
            {"name":"Stack diversity","score":3 if multi_stack else 0,"weight":3,"insight":"Multiple languages/tools referenced."},
        ])
    else:
        gh_sub.append({"name":"Public link present","score":0,"weight":3,"insight":"No GitHub link or username found."})

    # ---------------- LeetCode (0–9) ----------------
    lc_sub = []
    lc_score = 0
    if leetcode_presence:
        solved_hint = re.search(r'(\d{2,4})\s*\+?\s*(?:problems|questions|solutions)\b', text_lower)
        hard_hint   = has_word(resume_text, "hard", "dp", "graph", "greedy", "binary search", "segment tree", "fenwick")
        contest     = has_word(resume_text, "contest", "weekly", "biweekly", "ranking", "rating")
        baseline = 3
        if solved_hint:
            try:
                n = int(solved_hint.group(1))
                if n >= 300: baseline = 8
                elif n >= 200: baseline = 7
                elif n >= 100: baseline = 6
                elif n >= 50: baseline = 5
                else: baseline = 4
            except Exception:
                pass
        if hard_hint: baseline += 1
        if contest: baseline += 1
        lc_score = min(9, baseline)
        lc_sub = [
            {"name":"Profile presence","score":3,"weight":3,"insight":"LeetCode link or username found."},
            {"name":"Problem count","score":min(3, baseline-3),"weight":3,"insight":"Solved problems count inferred."},
            {"name":"Hard/contest","score":2 if (hard_hint or contest) else 0,"weight":3,"insight":"Advanced topics or contests mentioned."},
        ]
    else:
        lc_sub = [{"name":"Profile presence","score":0,"weight":3,"insight":"No LeetCode link or username found."}]

    # ---------------- LinkedIn (0–18) ----------------
    li_sub = []
    li_score = 0
    if linkedin_presence:
        headline = any(has_word(l.get("title","") or l.get("description",""), "engineer", "developer", "data", "software", "analyst", "scientist") for l in li_links)
        about    = any(has_word(l.get("description",""), "experience", "skills", "certification", "education", "project") for l in li_links)
        exp      = has_word(resume_text, "experience", "worked", "job", "position", "role", "internship")
        projects = has_word(resume_text, "project", "built", "developed", "created", "designed")
        skills   = has_word(resume_text, "skill", "proficient", "expert", "familiar", "knowledge")
        certs    = has_word(resume_text, "certification", "certified", "course", "training", "degree", "education")
        base = 5
        if headline: base += 2
        if about: base += 3
        if exp: base += 4
        if projects: base += 3
        if skills: base += 2
        if certs: base += 2
        li_score = min(18, base)
        li_sub = [
            {"name":"Headline","score":2 if headline else 0,"weight":2,"insight":"Professional headline present."},
            {"name":"About section","score":3 if about else 0,"weight":3,"insight":"About section or summary found."},
            {"name":"Experience","score":4 if exp else 0,"weight":4,"insight":"Experience section detected."},
            {"name":"Projects","score":3 if projects else 0,"weight":3,"insight":"Projects mentioned."},
            {"name":"Skills","score":2 if skills else 0,"weight":2,"insight":"Skills section found."},
            {"name":"Certs/Education","score":2 if certs else 0,"weight":2,"insight":"Certifications or education mentioned."},
        ]
    else:
        li_sub = [{"name":"Profile presence","score":0,"weight":2,"insight":"No LinkedIn link found."}]

    # ---------------- Portfolio (0–23) ----------------
    pf_sub = []
    pf_score = 0
    if portfolio_presence:
        pf_links = portfolio_links
        pf_count = len(pf_links)
        has_blog = len(blog_links) > 0
        has_notion = len(notion_links) > 0
        has_demo = len(demoish) > 0
        mentions_portfolio = has_word(resume_text, "portfolio", "personal website", "project", "demo")
        base = 5
        if pf_count == 1: base += 3
        elif pf_count >= 2: base += 6
        if has_blog: base += 3
        if has_notion: base += 2
        if has_demo: base += 4
        if mentions_portfolio: base += 3
        pf_score = min(23, base)
        pf_sub = [
            {"name":"Portfolio links","score":min(6, pf_count*3),"weight":6,"insight":f"{pf_count} portfolio link(s) found."},
            {"name":"Blog posts","score":3 if has_blog else 0,"weight":3,"insight":"Blog or articles found."},
            {"name":"Notion/docs","score":2 if has_notion else 0,"weight":2,"insight":"Notion or documentation links."},
            {"name":"Live demos","score":4 if has_demo else 0,"weight":4,"insight":"Live demo links found."},
            {"name":"Mentions in resume","score":3 if mentions_portfolio else 0,"weight":3,"insight":"Portfolio mentioned in resume."},
        ]
    else:
        pf_sub = [{"name":"Portfolio presence","score":0,"weight":6,"insight":"No portfolio links found."}]

    # ---------------- Resume ATS (0–23) ----------------
    ats_sub = []
    ats_score = 0
    if resume_text:
        kw_tech = sum(1 for k in ["python","sql","aws","docker","kubernetes","spark","airflow","javascript","typescript","java","go","rust","scala","kotlin","c++","c#","terraform","ansible","jenkins","ci/cd","git","github","gitlab","jira","agile","scrum","tableau","powerbi","excel","nosql","mongodb","postgresql","mysql","redis","kafka","rabbitmq","elasticsearch","snowflake","redshift","bigquery","databricks","hadoop","hive","hbase","cassandra","dynamodb","s3","ec2","lambda","glue","step functions","cloudformation","cloudwatch","x-ray","vpc","iam","route53","elb","nginx","apache","linux","bash","shell","powershell","nodejs","react","angular","vue","django","flask","fastapi","spring","express","graphql","rest","api","grpc","thrift","protobuf","avro","parquet","orc","csv","json","xml","yaml","toml","ini","env","dockerfile","docker-compose","k8s","helm","kustomize","argo","tekton","spinnaker","jenkins","circleci","github actions","gitlab ci","travis","teamcity","bamboo","ansible","puppet","chef","saltstack","terraform","cloudformation","pulumi","cdk","serverless","sam","chalice","zappa","vercel","netlify","heroku","firebase","supabase","auth0","okta","cognito","oauth","jwt","ssl","tls","https","ssh","vpn","ipsec","wireguard","openvpn","ldap","kerberos","saml","oidc","mfa","2fa","totp","hotp","webauthn","fido","u2f","biometrics","facial recognition","iris scan","fingerprint","voice recognition","behavioral biometrics","ai","ml","machine learning","deep learning","neural networks","cnn","rnn","lstm","gru","transformer","bert","gpt","t5","vit","resnet","inception","efficientnet","mobilenet","yolo","ssd","faster r-cnn","mask r-cnn","retinanet","centernet","detr","pointpillars","pointrcnn","second","pv-rcnn","part-a2","3dssd","voxelnet","pointnet","pointnet++","pointcnn","dgcnn","kpconv","randla-net","polarnet","spconv","minkowskiengine","torchsparse","openpcdet","mmdetection3d","detectron2","mmdetection","yolov5","yolov6","yolov7","yolov8","yolov9","yolox","scaled-yolov4","efficientdet","nanodet","pp-yolo","ppyoloe","ppyolov2","centernet","cornernet","fcos","atss","gfl","dynamicrpn","reppoints","foveabox","freenet","sparse r-cnn","querydet","detic","deformable detr","conditional detr","dab-detr","dn-detr","dino","mask2former","maskformer","k-net","max-deeplab","maskclip","mask dino","groupvit","x-decoder","open-seed","sam","segment anything","fastsam","efficientsam","mobile sam","edge sam","sam2","grounding dino","grounding sam","owl","owlv2","glip","glipv2","detclip","grit","uni-detr","uniperceptor","unidet","ovd","open-vocabulary detection","zero-shot detection","few-shot detection","semi-supervised detection","weakly-supervised detection","unsupervised detection","self-supervised detection","contrastive learning","momentum contrast","moco","simclr","byol","swav","barlow twins","vicreg","dino","ibot","mae","masked autoencoder","simsiam","nnclr","supcon","deepcluster","scan","sela","pcl","cpc","amdim","bigbigan","stylegan","stylegan2","stylegan3","progan","sggan","logan","wgan","wgan-gp","lsgan","rsgan","ragan","hingegan","loss sensitive gan","ebgan","began","margin gan","f-gan","mmgan","nsgan","sngan","sagan","biggan","bigbigan","trgan","stylegan-ada","stylegan2-ada","stylegan3-ada","training","fine-tuning","transfer learning","domain adaptation","domain generalization","test-time adaptation","test-time training","meta-learning","few-shot learning","zero-shot learning","multi-task learning","continual learning","lifelong learning","online learning","active learning","semi-supervised learning","weakly-supervised learning","self-supervised learning","unsupervised learning","reinforcement learning","imitation learning","inverse reinforcement learning","offline rl","online rl","batch rl","model-based rl","model-free rl","policy gradient","actor-critic","q-learning","sarsa","dqn","ddpg","td3","sac","ppo","trpo","mpc","ilqr","ddim","ddpm","score-based","diffusion","normalizing flows","real nvp","glow","maf","iaf","nice","ffjord","sde","ode","neural ode","hamiltonian nn","lagrangian nn","symplectic nn","geometric nn","graph nn","gnn","gcn","gat","graphsage","gin","pna","mpnn","transformers","attention","self-attention","multi-head attention","transformer","bert","gpt","t5","vit","swin","deit","cait","crossvit","levit","convit","t2t-vit","pit","xcit","coat","cvt","twins","pvt","shuffle transformer","mobilevit","edgevit","efficientformer","mobileformer","poolformer","uniformer","linformer","performer","nyströmformer","longformer","bigbird","reformer","linformer","sinkhorn transformer","rfa","linear transformer","synthesizer","rfa","fast attention","flash attention","block-sparse attention","longshort-transformer","et","informer","autoformer","fedformer","stationary","non-stationary","time series","forecasting","anomaly detection","change point detection","event detection","segmentation","classification","regression","clustering","dimensionality reduction","feature selection","feature extraction","manifold learning","pca","ica","lda","nmf","tsne","umap","phate","diffusion maps","isomap","lle","mds","spectral embedding","kernel pca","autoencoder","vae","cvae","vq-vae","vq-vae-2","nva","diffusion autoencoder","gan","vae-gan","cyclegan","unit","munit","stargan","stargan2","ganilla","attentiongan","stylegan","biggan","stylegan-ada","stylegan2-ada","stylegan3-ada","training","fine-tuning","transfer learning","domain adaptation","domain generalization","test-time adaptation","test-time training","meta-learning","few-shot learning","zero-shot learning","multi-task learning","continual learning","lifelong learning","online learning","active learning","semi-supervised learning","weakly-supervised learning","self-supervised learning","unsupervised learning","reinforcement learning","imitation learning","inverse reinforcement learning","offline rl","online rl","batch rl","model-based rl","model-free rl","policy gradient","actor-critic","q-learning","sarsa","dqn","ddpg","td3","sac","ppo","trpo","mpc","ilqr","ddim","ddpm","score-based","diffusion","normalizing flows","real nvp","glow","maf","iaf","nice","ffjord","sde","ode","neural ode","hamiltonian nn","lagrangian nn","symplectic nn","geometric nn","graph nn","gnn","gcn","gat","graphsage","gin","pna","mpnn","transformers","attention","self-attention","multi-head attention","transformer","bert","gpt","t5","vit","swin","deit","cait","crossvit","levit","convit","t2t-vit","pit","xcit","coat","cvt","twins","pvt","shuffle transformer","mobilevit","edgevit","efficientformer","mobileformer","poolformer","uniformer","linformer","performer","nyströmformer","longformer","bigbird","reformer","linformer","sinkhorn transformer","rfa","linear transformer","synthesizer","rfa","fast attention","flash attention","block-sparse attention","longshort-transformer","et","informer","autoformer","fedformer","stationary","non-stationary","time series","forecasting","anomaly detection","change point detection","event detection","segmentation","classification","regression","clustering","dimensionality reduction","feature selection","feature extraction","manifold learning","pca","ica","lda","nmf","tsne","umap","phate","diffusion maps","isomap","lle","mds","spectral embedding","kernel pca","autoencoder","vae","cvae","vq-vae","vq-vae-2","nva","diffusion autoencoder","gan","vae-gan","cyclegan","unit","munit","stargan","stargan2","ganilla","attentiongan","stylegan","biggan","stylegan-ada","stylegan2-ada","stylegan3-ada"] if k in text_lower)
        kw_soft = sum(1 for k in ["leadership","team","communication","problem solving","critical thinking","creativity","adaptability","time management","project management","agile","scrum","kanban","lean","devops","mlops","dataops","gitops","finops","secops","aiops","modelops","feature store","mlflow","kubeflow","metaflow","polyaxon","sagemaker","vertex ai","azure ml","databricks","hugging face","wandb","comet","neptune","determined","pachyderm","dvc","lakefs","delta lake","iceberg","hudi","dremio","presto","trino","athena","bigquery","snowflake","redshift","databricks","synapse","firebolt","motherduck","clickhouse","druid","pinot","kylin","greenplum","vertica","teradata","netezza","exadata","oracle","sql server","mysql","postgresql","sqlite","cockroachdb","yugabyte","tidb","cassandra","scylla","hbase","accumulo","bigtable","dynamodb","cosmos db","documentdb","mongodb","couchbase","redis","memcached","etcd","consul","zookeeper","kafka","pulsar","rabbitmq","activemq","nats","jetstream","nsq","beanstalkd","gearman","celery","dramatiq","rq","huey","apscheduler","airflow","prefect","dagster","luigi","argo","tekton","jenkins","circleci","github actions","gitlab ci","travis","teamcity","bamboo","spinnaker","go cd","concourse","bazel","buck","pants","please","soong","make","cmake","ninja","meson","scons","ant","maven","gradle","sbt","cargo","npm","yarn","pnpm","pip","pipenv","poetry","conda","mamba","spack","nix","guix","docker","podman","containerd","cri-o","rkt","lxc","lxd","systemd-nspawn","firecracker","gvisor","kata","nabla","unikernel","wasm","wasmer","wasmtime","wasi","emscripten","cheerp","j2cl","bytenode","napa","node","deno","bun","quickjs","hermes","v8","javascriptcore","spidermonkey","graalvm","truffle","substrate","ink","solang","solidity","vyper","yul","lll","serpent","mutan","bamboo","scilla","lity","plutus","marlowe","glow","reach","act","cairo","noir","leo","move","sui move","aptos move","fuel","sway","rust","go","zig","nim","crystal","v","odin","jai","hare","carbon","cpp2","circle","dagger","vala","genie","seed7","chapel","x10","fortress","clojure","racket","scheme","common lisp","elisp","hy","janet","fennel","lfe","picolisp","newlisp","guile","chicken","racket","chez","gambit","gerbil","picrin","sagittarius","cyclone","owl lisp","ulisp","esp lisp","microlisp","picolisp","newlisp","clojurescript","lumen","joker","balisong","lux","mal","make a lisp","lisp in a weekend","lisp in a day","lisp in an hour","lisp in a minute","lisp in a second","lisp in a millisecond","lisp in a microsecond","lisp in a nanosecond","lisp in a picosecond","lisp in a femtosecond","lisp in a attosecond","lisp in a zeptosecond","lisp in a yoctosecond","lisp in a planck time","lisp in a jiffy","lisp in a shake","lisp in a moment","lisp in a while","lisp in a bit","lisp in a tick","lisp in a tock","lisp in a flash","lisp in a blink","lisp in a snap","lisp in a trice","lisp in a twinkling","lisp in a jiffy","lisp in a mo","lisp in a sec","lisp in a min","lisp in a hr","lisp in a day","lisp in a wk","lisp in a mo","lisp in a yr","lisp in a decade","lisp in a century","lisp in a millennium","lisp in a eon","lisp in a era","lisp in a period","lisp in a epoch","lisp in a age","lisp in a stage","lisp in a phase","lisp in a step","lisp in a level","lisp in a tier","lisp in a layer","lisp in a stratum","lisp in a grade","lisp in a rank","lisp in a class","lisp in a order","lisp in a family","lisp in a genus","lisp in a species","lisp in a variety","lisp in a form","lisp in a type","lisp in a kind","lisp in a sort","lisp in a manner","lisp in a style","lisp in a fashion","lisp in a way","lisp in a method","lisp in a mode","lisp in a system","lisp in a technique","lisp in a procedure","lisp in a process","lisp in a operation","lisp in a action","lisp in a activity","lisp in a task","lisp in a job","lisp in a work","lisp in a labor","lisp in a effort","lisp in a exertion","lisp in a strain","lisp in a stress","lisp in a pressure","lisp in a tension","lisp in a force","lisp in a power","lisp in a energy","lisp in a strength","lisp in a might","lisp in a potency","lisp in a capacity","lisp in a ability","lisp in a capability","lisp in a competence","lisp in a proficiency","lisp in a skill","lisp in a talent","lisp in a gift","lisp in a flair","lisp in a knack","lisp in a bent","lisp in a aptitude","lisp in a facility","lisp in a dexterity","lisp in a adroitness","lisp in a cleverness","lisp in a ingenuity","lisp in a inventiveness","lisp in a creativity","lisp in a originality","lisp in a imagination","lisp in a inspiration","lisp in a vision","lisp in a insight","lisp in a perception","lisp in a discernment","lisp in a judgment","lisp in a wisdom","lisp in a knowledge","lisp in a understanding","lisp in a comprehension","lisp in a grasp","lisp in a mastery","lisp in a command","lisp in a control","lisp in a dominion","lisp in a sovereignty","lisp in a supremacy","lisp in a authority","lisp in a influence","lisp in a sway","lisp in a clout","lisp in a pull","lisp in a weight","lisp in a importance","lisp in a significance","lisp in a consequence","lisp in a moment","lisp in a value","lisp in a worth","lisp in a merit","lisp in a virtue","lisp in a excellence","lisp in a quality","lisp in a property","lisp in a attribute","lisp in a characteristic","lisp in a feature","lisp in a trait","lisp in a aspect","lisp in a facet","lisp in a side","lisp in a angle","lisp in a perspective","lisp in a viewpoint","lisp in a standpoint","lisp in a position","lisp in a stance","lisp in a attitude","lisp in a disposition","lisp in a temperament","lisp in a mood","lisp in a humor","lisp in a spirit","lisp in a sentiment","lisp in a feeling","lisp in a emotion","lisp in a passion","lisp in a desire","lisp in a wish","lisp in a want","lisp in a need","lisp in a craving","lisp in a longing","lisp in a yearning","lisp in a hankering","lisp in a itch","lisp in a urge","lisp in a impulse","lisp in a drive","lisp in a motive","lisp in a incentive","lisp in a inducement","lisp in a stimulus","lisp in a spur","lisp in a goad","lisp in a prod","lisp in a push","lisp in a pressure","lisp in a force","lisp in a compulsion","lisp in a constraint","lisp in a obligation","lisp in a duty","lisp in a responsibility","lisp in a charge","lisp in a task","lisp in a assignment","lisp in a mission","lisp in a function","lisp in a role","lisp in a part","lisp in a capacity","lisp in a office","lisp in a post","lisp in a position","lisp in a station","lisp in a status","lisp in a rank","lisp in a grade","lisp in a class","lisp in a order","lisp in a category","lisp in a group","lisp in a set","lisp in a collection","lisp in a assortment","lisp in a variety","lisp in a diversity","lisp in a mixture","lisp in a blend","lisp in a combination","lisp in a compound","lisp in a composite","lisp in a synthesis","lisp in a amalgamation","lisp in a fusion","lisp in a union","lisp in a junction","lisp in a connection","lisp in a link","lisp in a bond","lisp in a tie","lisp in a relationship","lisp in a association","lisp in a affiliation","lisp in a alliance","lisp in a partnership","lisp in a cooperation","lisp in a collaboration","lisp in a teamwork","lisp in a coordination","lisp in a harmony","lisp in a concord","lisp in a agreement","lisp in a consensus","lisp in a unity","lisp in a solidarity","lisp in a community","lisp in a society","lisp in a fellowship","lisp in a brotherhood","lisp in a sisterhood","lisp in a fraternity","lisp in a sorority","lisp in a club","lisp in a group","lisp in a team","lisp in a squad","lisp in a crew","lisp in a gang","lisp in a band","lisp in a troupe","lisp in a company","lisp in a corporation","lisp in a firm","lisp in a business","lisp in a enterprise","lisp in a venture","lisp in a undertaking","lisp in a project","lisp in a scheme","lisp in a plan","lisp in a program","lisp in a campaign","lisp in a operation","lisp in a activity","lisp in a action","lisp in a move","lisp in a step","lisp in a measure","lisp in a initiative","lisp in a effort","lisp in a endeavor","lisp in a strive","lisp in a struggle","lisp in a fight","lisp in a battle","lisp in a war","lisp in a conflict","lisp in a contest","lisp in a competition","lisp in a rivalry","lisp in a race","lisp in a chase","lisp in a pursuit","lisp in a quest","lisp in a search","lisp in a hunt","lisp in a exploration","lisp in a investigation","lisp in a inquiry","lisp in a research","lisp in a study","lisp in a examination","lisp in a analysis","lisp in a scrutiny","lisp in a inspection","lisp in a review","lisp in a survey","lisp in a poll","lisp in a census","lisp in a count","lisp in a tally","lisp in a calculation","lisp in a computation","lisp in a estimation","lisp in a assessment","lisp in a evaluation","lisp in a appraisal","lisp in a judgment","lisp in a opinion","lisp in a view","lisp in a belief","lisp in a conviction","lisp in a faith","lisp in a trust","lisp in a confidence","lisp in a reliance","lisp in a dependence","lisp in a hope","lisp in a expectation","lisp in a anticipation","lisp in a prospect","lisp in a outlook","lisp in a forecast","lisp in a prediction","lisp in a prophecy","lisp in a divination","lisp in a augury","lisp in a omen","lisp in a sign","lisp in a signal","lisp in a indicator","lisp in a index","lisp in a gauge","lisp in a measure","lisp in a standard","lisp in a criterion","lisp in a benchmark","lisp in a yardstick","lisp in a touchstone","lisp in a test","lisp in a trial","lisp in a experiment","lisp in a proof","lisp in a evidence","lisp in a testimony","lisp in a witness","lisp in a attestation","lisp in a confirmation","lisp in a verification","lisp in a validation","lisp in a authentication","lisp in a certification","lisp in a accreditation","lisp in a endorsement","lisp in a approval","lisp in a sanction","lisp in a authorization","lisp in a license","lisp in a permit","lisp in a warrant","lisp in a charter","lisp in a franchise","lisp in a privilege","lisp in a right","lisp in a entitlement","lisp in a claim","lisp in a demand","lisp in a request","lisp in a requirement","lisp in a necessity","lisp in a need","lisp in a want","lisp in a desire","lisp in a wish","lisp in a longing","lisp in a yearning","lisp in a hankering","lisp in a craving","lisp in a appetite","lisp in a hunger","lisp in a thirst","lisp in a lust","lisp in a passion","lisp in a urge","lisp in a impulse","lisp in a drive","lisp in a motive","lisp in a incentive","lisp in a inducement","lisp in a stimulus","lisp in a spur","lisp in a goad","lisp in a prod","lisp in a push","lisp in a pressure","lisp in a force","lisp in a compulsion","lisp in a constraint","lisp in a obligation","lisp in a duty","lisp in a responsibility","lisp in a charge","lisp in a task","lisp in a assignment","lisp in a mission","lisp in a function","lisp in a role","lisp in a part","lisp in a capacity","lisp in a office","lisp in a post","lisp in a position","lisp in a station","lisp in a status","lisp in a rank","lisp in a grade","lisp in a class","lisp in a order","lisp in a category","lisp in a group","lisp in a set","lisp in a collection","lisp in a assortment","lisp in a variety","lisp in a diversity","lisp in a mixture","lisp in a blend","lisp in a combination","lisp in a compound","lisp in a composite","lisp in a synthesis","lisp in a amalgamation","lisp in a fusion","lisp in a union","lisp in a junction","lisp in a connection","lisp in a link","lisp in a bond","lisp in a tie","lisp in a relationship","lisp in a association","lisp in a affiliation","lisp in a alliance","lisp in a partnership","lisp in a cooperation","lisp in a collaboration","lisp in a teamwork","lisp in a coordination","lisp in a harmony","lisp in a concord","lisp in a agreement","lisp in a consensus","lisp in a unity","lisp in a solidarity","lisp in a community","lisp in a society","lisp in a fellowship","lisp in a brotherhood","lisp in a sisterhood","lisp in a fraternity","lisp in a sorority","lisp in a club","lisp in a group","lisp in a team","lisp in a squad","lisp in a crew","lisp in a gang","lisp in a band","lisp in a troupe","lisp in a company","lisp in a corporation","lisp in a firm","lisp in a business","lisp in a enterprise","lisp in a venture","lisp in a undertaking","lisp in a project","lisp in a scheme","lisp in a plan","lisp in a program","lisp in a campaign","lisp in a operation","lisp in a activity","lisp in a action","lisp in a move","lisp in a step","lisp in a measure","lisp in a initiative","lisp in a effort","lisp in a endeavor","lisp in a strive","lisp in a struggle","lisp in a fight","lisp in a battle","lisp in a war","lisp in a conflict","lisp in a contest","lisp in a competition","lisp in a rivalry","lisp in a race","lisp in a chase","lisp in a pursuit","lisp in a quest","lisp in a search","lisp in a hunt","lisp in a exploration","lisp in a investigation","lisp in a inquiry","lisp in a research","lisp in a study","lisp in a examination","lisp in a analysis","lisp in a scrutiny","lisp in a inspection","lisp in a review","lisp in a survey","lisp in a poll","lisp in a census","lisp in a count","lisp in a tally","lisp in a calculation","lisp in a computation","lisp in a estimation","lisp in a assessment","lisp in a evaluation","lisp in a appraisal","lisp in a judgment","lisp in a opinion","lisp in a view","lisp in a belief","lisp in a conviction","lisp in a faith","lisp in a trust","lisp in a confidence","lisp in a reliance","lisp in a dependence","lisp in a hope","lisp in a expectation","lisp in a anticipation","lisp in a prospect","lisp in a outlook","lisp in a forecast","lisp in a prediction","lisp in a prophecy","lisp in a divination","lisp in a augury","lisp in a omen","lisp in a sign","lisp in a signal","lisp in a indicator","lisp in a index","lisp in a gauge","lisp in a measure","lisp in a standard","lisp in a criterion","lisp in a benchmark","lisp in a yardstick","lisp in a touchstone","lisp in a test","lisp in a trial","lisp in a experiment","lisp in a proof","lisp in a evidence","lisp in a testimony","lisp in a witness","lisp in a attestation","lisp in a confirmation","lisp in a verification","lisp in a validation","lisp in a authentication","lisp in a certification","lisp in a accreditation","lisp in a endorsement","lisp in a approval","lisp in a sanction","lisp in a authorization","lisp in a license","lisp in a permit","lisp in a warrant","lisp in a charter","lisp in a franchise","lisp in a privilege","lisp in a right","lisp in a entitlement","lisp in a claim","lisp in a demand","lisp in a request","lisp in a requirement","lisp in a necessity","lisp in a need","lisp in a want","lisp in a desire","lisp in a wish","lisp in a longing","lisp in a yearning","lisp in a hankering","lisp in a craving","lisp in a appetite","lisp in a hunger","lisp in a thirst","lisp in a lust","lisp in a passion","lisp in a urge","lisp in a impulse","lisp in a drive","lisp in a motive","lisp in a incentive","lisp in a inducement","lisp in a stimulus","lisp in a spur","lisp in a goad","lisp in a prod","lisp in a push","lisp in a pressure","lisp in a force","lisp in a compulsion","lisp in a constraint","lisp in a obligation","lisp in a duty","lisp in a responsibility","lisp in a charge","lisp in a task","lisp in a assignment","lisp in a mission","lisp in a function","lisp in a role","lisp in a part","lisp in a capacity","lisp in a office","lisp in a post","lisp in a position","lisp in a station","lisp in a status","lisp in a rank","lisp in a grade","lisp in a class","lisp in a order","lisp in a category","lisp in a group","lisp in a set","lisp in a collection","lisp in a assortment","lisp in a variety","lisp in a diversity","lisp in a mixture","lisp in a blend","lisp in a combination","lisp in a compound","lisp in a composite","lisp in a synthesis","lisp in a amalgamation","lisp in a fusion","lisp in a union","lisp in a junction","lisp in a connection","lisp in a link","lisp in a bond","lisp in a tie","lisp in a relationship","lisp in a association","lisp in a affiliation","lisp in a alliance","lisp in a partnership","lisp in a cooperation","lisp in a collaboration","lisp in a teamwork","lisp in a coordination","lisp in a harmony","lisp in a concord","lisp in a agreement","lisp in a consensus","lisp in a unity","lisp in a solidarity","lisp in a community","lisp in a society","lisp in a fellowship","lisp in a brotherhood","lisp in a sisterhood","lisp in a fraternity","lisp in a sorority","lisp in a club","lisp in a group","lisp in a team","lisp in a squad","lisp in a crew","lisp in a gang","lisp in a band","lisp in a troupe","lisp in a company","lisp in a corporation","lisp in a firm","lisp in a business","lisp in a enterprise","lisp in a venture","lisp in a undertaking","lisp in a project","lisp in a scheme","lisp in a plan","lisp in a program","lisp in a campaign","lisp in a operation","lisp in a activity","lisp in a action","lisp in a move","lisp in a step","lisp in a measure","lisp in a initiative","lisp in a effort","lisp in a endeavor","lisp in a strive","lisp in a struggle","lisp in a fight","lisp in a battle","lisp in a war","lisp in a conflict","lisp in a contest","lisp in a competition","lisp in a rivalry","lisp in a race","lisp in a chase","lisp in a pursuit","lisp in a quest","lisp in a search","lisp in a hunt","lisp in a exploration","lisp in a investigation","lisp in a inquiry","lisp in a research","lisp in a study","lisp in a examination","lisp in a analysis","lisp in a scrutiny","lisp in a inspection","lisp in a review","lisp in a survey","lisp in a poll","lisp in a census","lisp in a count","lisp in a tally","lisp in a calculation","lisp in a computation","lisp in a estimation","lisp in a assessment","lisp in a evaluation","lisp in a appraisal","lisp in a judgment","lisp in a opinion","lisp in a view","lisp in a belief","lisp in a conviction","lisp in a faith","lisp in a trust","lisp in a confidence","lisp in a reliance","lisp in a dependence","lisp in a hope","lisp in a expectation","lisp in a anticipation","lisp in a prospect","lisp in a outlook","lisp in a forecast","lisp in a prediction","lisp in a prophecy","lisp in a divination","lisp in a augury","lisp in a omen","lisp in a sign","lisp in a signal","lisp in a indicator","lisp in a index","lisp in a gauge","lisp in a measure","lisp in a standard","lisp in a criterion","lisp in a benchmark","lisp in a yardstick","lisp in a touchstone","lisp in a test","lisp in a trial","lisp in a experiment","lisp in a proof","lisp in a evidence","lisp in a testimony","lisp in a witness","lisp in a attestation","lisp in a confirmation","lisp in a verification","lisp in a validation","lisp in a authentication","lisp in a certification","lisp in a accreditation","lisp in a endorsement","lisp in a approval","lisp in a sanction","lisp in a authorization","lisp in a license","lisp in a permit","lisp in a warrant","lisp in a charter","lisp in a franchise","lisp in a privilege","lisp in a right","lisp in a entitlement","lisp in a claim","lisp in a demand","lisp in a request","lisp in a requirement","lisp in a necessity","lisp in a need","lisp in a want","lisp in a desire","lisp in a wish","lisp in a longing","lisp in a yearning","lisp in a hankering","lisp in a craving","lisp in a appetite","lisp in a hunger","lisp in a thirst","lisp in a lust","lisp in a passion","lisp in a urge","lisp in a impulse","lisp in a drive","lisp in a motive","lisp in a incentive","lisp in a inducement","lisp in a stimulus","lisp in a spur","lisp in a goad","lisp in a prod","lisp in a push","lisp in a pressure","lisp in a force","lisp in a compulsion","lisp in a constraint","lisp in a obligation","lisp in a duty","lisp in a responsibility","lisp in a charge","lisp in a task","lisp in a assignment","lisp in a mission","lisp in a function","lisp in a role","lisp in a part","lisp in a capacity","lisp in a office","lisp in a post","lisp in a position","lisp in a station","lisp in a status","lisp in a rank","lisp in a grade","lisp in a class","lisp in a order","lisp in a category","lisp in a group","lisp in a set","lisp in a collection","lisp in a assortment","lisp in a variety","lisp in a diversity","lisp in a mixture","lisp in a blend","lisp in a combination","lisp in a compound","lisp in a composite","lisp in a synthesis","lisp in a amalgamation","lisp in a fusion","lisp in a union","lisp in a junction","lisp in a connection","lisp in a link","lisp in a bond","lisp in a tie","lisp in a relationship","lisp in a association","lisp in a affiliation","lisp in a alliance","lisp in a partnership","lisp in a cooperation","lisp in a collaboration","lisp in a teamwork","lisp in a coordination","lisp in a harmony","lisp in a concord","lisp in a agreement","lisp in a consensus","lisp in a unity","lisp in a solidarity","lisp in a community","lisp in a society","lisp in a fellowship","lisp in a brotherhood","lisp in a sisterhood","lisp in a fraternity","lisp in a sorority","lisp in a club","lisp in a group","lisp in a team","lisp in a squad","lisp in a crew","lisp in a gang","lisp in a band","lisp in a troupe","lisp in a company","lisp in a corporation","lisp in a firm","lisp in a business","lisp in a enterprise","lisp in a venture","lisp in a undertaking","lisp in a project","lisp in a scheme","lisp in a plan","lisp in a program","lisp in a campaign","lisp in a operation","lisp in a activity","lisp in a action","lisp in a move","lisp in a step","lisp in a measure","lisp in a initiative","lisp in a effort","lisp in a endeavor","lisp in a strive","lisp in a struggle","lisp in a fight","lisp in a battle","lisp in a war","lisp in a conflict","lisp in a contest","lisp in a competition","lisp in a rivalry","lisp in a race","lisp in a chase","lisp in a pursuit","lisp in a quest","lisp in a search","lisp in a hunt","lisp in a exploration","lisp in a investigation","lisp in a inquiry","lisp in a research","lisp in a study","lisp in a examination","lisp in a analysis","lisp in a scrutiny","lisp in a inspection","lisp in a review","lisp in a survey","lisp in a poll","lisp in a census","lisp in a count","lisp in a tally","lisp in a calculation","lisp in a computation","lisp in a estimation","lisp in a assessment","lisp in a evaluation","lisp in a appraisal","lisp in a judgment","lisp in a opinion","lisp in a view","lisp in a belief","lisp in a conviction","lisp in a faith","lisp in a trust","lisp in a confidence","lisp in a reliance","lisp in a dependence","lisp in a hope","lisp in a expectation","lisp in a anticipation","lisp in a prospect","lisp in a outlook","lisp in a forecast","lisp in a prediction","lisp in a prophecy","lisp in a divination","lisp in a augury","lisp in a omen","lisp in a sign","lisp in a signal","lisp in a indicator","lisp in a index","lisp in a gauge","lisp in a measure","lisp in a standard","lisp in a criterion","lisp in a benchmark","lisp in a yardstick","lisp in a touchstone","lisp in a test","lisp in a trial","lisp in a experiment","lisp in a proof","lisp in a evidence","lisp in a testimony","lisp in a witness","lisp in a attestation","lisp in a confirmation","lisp in a verification","lisp in a validation","lisp in a authentication","lisp in a certification","lisp in a accreditation","lisp in a endorsement","lisp in a approval","lisp in a sanction","lisp in a authorization","lisp in a license","lisp in a permit","lisp in a warrant","lisp in a charter","lisp in a franchise","lisp in a privilege","lisp in a right","lisp in a entitlement","lisp in a claim","lisp in a demand","lisp in a request","lisp in a requirement","lisp in a necessity","lisp in a need","lisp in a want","lisp in a desire","lisp in a wish","lisp in a longing","lisp in a yearning","lisp in a hankering","lisp in a craving","lisp in a appetite","lisp in a hunger","lisp in a thirst","lisp in a lust","lisp in a passion","lisp in a urge","lisp in a impulse","lisp in a drive","lisp in a motive","lisp in a incentive","lisp in a inducement","lisp in a stimulus","lisp in a spur","lisp in a goad","lisp in a prod","lisp in a push","lisp in a pressure","lisp in a force","lisp in a compulsion","lisp in a constraint","lisp in a obligation","lisp in a duty","lisp in a responsibility","lisp in a charge","lisp in a task","lisp in a assignment","lisp in a mission","lisp in a function","lisp in a role","lisp in a part","lisp in a capacity","lisp in a office","lisp in a post","lisp in a position","lisp in a station","lisp in a status","lisp in a rank","lisp in a grade","lisp in a class","lisp in a order","lisp in a category","lisp in a group","lisp in a set","lisp in a collection","lisp in a assortment","lisp in a variety","lisp in a diversity","lisp in a mixture","lisp in a blend","lisp in a combination","lisp in a compound","lisp in a composite","lisp in a synthesis","lisp in a amalgamation","lisp in a fusion","lisp in a union","lisp in a junction","lisp in a connection","lisp in a link","lisp in a bond","lisp in a tie","lisp in a relationship","lisp in a association","lisp in a affiliation","lisp in a alliance","lisp in a partnership","lisp in a cooperation","lisp in a collaboration","lisp in a teamwork","lisp in a coordination","lisp in a harmony","lisp in a concord","lisp in a agreement","lisp极","lisp in a consensus","lisp in a unity","lisp in a solidarity","lisp in a community","lisp in a society","lisp in a fellowship","lisp in a brotherhood","lisp in a sisterhood","lisp in a fraternity","lisp in a sorority","lisp in a club","极"] if k in text_lower)
        base = 5
        if kw_tech >= 20: base += 5
        elif kw_tech >= 10: base += 3
        elif kw_tech >= 5: base += 1
        if kw_soft >= 5: base += 3
        elif kw_soft >= 3: base += 1
        has_quant = bool(re.search(r'\b\d+%|\$\d+|\d+\s*(?:years?|months?|weeks?|days?|hours?|minutes?|seconds?)\b', resume_text))
        if has_quant: base += 2
        has_contact = bool(re.search(r'\b(?:phone|email|@|\.com|\.org|\.net|linkedin\.com|github\.com|leetcode\.com)\b', resume_text, re.I))
        if has_contact: base += 2
        has_edu = bool(re.search(r'\b(?:b\.?a|b\.?s|b\.?eng|m\.?a|m\.?s|m\.?eng|ph\.?d|doctorate|master|bachelor|diploma|degree|certificate|certification)\b', resume_text, re.I))
        if has_edu: base += 2
        ats_score = min(23, base)
        ats_sub = [
            {"name":"Technical keywords","score":min(5, kw_tech//4),"weight":5,"insight":f"{kw_tech} technical terms found."},
            {"name":"Soft skills","score":min(3, kw_soft//2),"weight":3,"insight":f"{kw_soft} soft skill terms found."},
            {"name":"Quantified results","score":2 if has_quant else 0,"weight":2,"insight":"Quantified achievements present."},
            {"name":"Contact info","score":2 if has_contact else 0,"weight":2,"insight":"Contact details included."},
            {"name":"Education","score":2 if has_edu else 0,"weight":2,"insight":"Education section found."},
        ]
    else:
        ats_sub = [{"name":"Resume text","score":0,"weight":5,"insight":"No resume text provided."}]

    # ---------------- Certifications & Branding (0–9) ----------------
    cert_sub = []
    cert_score = 0
    if cert_presence:
        cert_count = len(re.findall(r'\b(certification|certified|certificate)\b', resume_text or "", re.I))
        has_brand = has_word(resume_text, "personal brand", "branding", "online presence", "thought leadership")
        has_blog = len(blog_links) > 0
        base = 3
        if cert_count >= 3: base += 3
        elif cert_count >= 1: base += 2
        if has_brand: base += 2
        if has_blog: base += 1
        cert_score = min(9, base)
        cert_sub = [
            {"name":"Certifications","score":min(3, cert_count),"weight":3,"insight":f"{cert_count} certification(s) mentioned."},
            {"name":"Personal branding","score":2 if has_brand else 0,"weight":2,"insight":"Personal branding mentioned."},
            {"name":"Blog/articles","score":1 if has_blog else 0,"weight":1,"insight":"Blog or articles found."},
        ]
    else:
        cert_sub = [{"name":"Certifications","score":0,"weight":3,"insight":"No certifications found."}]

    # ---------------- Overall ----
    sections = {
        "GitHub Profile": {
            "score": gh_score,
            "grade": grade_tag(gh_score, 27),
            "weight": TECHNICAL_WEIGHTS["GitHub Profile"],
            "sub_criteria": gh_sub
        },
        "LinkedIn": {
            "score": li_score,
            "grade": grade_tag(li_score, 18),
            "weight": TECHNICAL_WEIGHTS["LinkedIn"],
            "sub_criteria": li_sub
        },
        "Portfolio Website": {
            "score": pf_score,
            "grade": grade_tag(pf_score, 23),
            "weight": TECHNICAL_WEIGHTS["Portfolio Website"],
            "sub_criteria": pf_sub
        },
        "Resume (ATS Score)": {
            "score": ats_score,
            "grade": grade_tag(ats_score, 23),
            "weight": TECHNICAL_WEIGHTS["Resume (ATS Score)"],
            "sub_criteria": ats_sub
        },
        "Certifications & Branding": {
            "score": cert_score,
            "grade": grade_tag(cert_score, 9),
            "weight": TECHNICAL_WEIGHTS["Certifications & Branding"],
            "sub_criteria": cert_sub
        },
        "LeetCode/DSA Skills": {
            "score": lc_score,
            "grade": grade_tag(lc_score, 9),
            "weight": TECHNICAL_WEIGHTS["LeetCode/DSA Skills"],
            "sub_criteria": lc_sub
        }
    }

    total_weighted = sum(s["score"] for s in sections.values())
    total_max = sum(TECHNICAL_WEIGHTS.values())
    overall_score = round((total_weighted / total_max) * 100, 1) if total_max > 0 else 0
    overall_grade = get_grade_tag(overall_score)

    suggestions = []
    if gh_score < 10:
        suggestions.append("Improve GitHub: add more projects, READMEs, tests, or live demos.")
    if li_score < 6:
        suggestions.append("Enhance LinkedIn: complete profile with experience, skills, and projects.")
    if pf_score < 10:
        suggestions.append("Build portfolio: create personal website, blog, or project showcases.")
    if ats_score < 10:
        suggestions.append("Optimize resume: add more keywords, quantify results, and include contact info.")
    if cert_score < 3:
        suggestions.append("Consider certifications: they can boost credibility and ATS scores.")
    if lc_score < 3:
        suggestions.append("Practice DSA: LeetCode or similar platforms can demonstrate problem-solving skills.")

    return {
        "sections": sections,
        "overall_score_average": overall_score,
        "overall_grade": overall_grade,
        "suggestions": suggestions
    }

# --- Plotting functions (no change needed for library version) ---
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import io
import base64
from .utils import get_grade_tag # Import from local utils for consistency

def prepare_chart_data(score_breakdown):
    """
    Prepares a dictionary of data that can be used directly by Chart.js.
    """
    labels = list(score_breakdown.keys())
    scores = [data['score'] for data in score_breakdown.values()]
    
    chart_colors = []
    for data in score_breakdown.values():
        grade = data['grade'].lower()
        if grade == 'excellent':
            chart_colors.append('#4CAF50')
        elif grade == 'good':
            chart_colors.append('#2196F3')
        elif grade == 'average':
            chart_colors.append('#FF9800')
        else:
            chart_colors.append('#dc3545')
    
    return {
        "labels": labels,
        "scores": scores,
        "backgroundColors": chart_colors,
    }

def generate_pie_chart(sections):
    """
    Generates a pie chart from section scores and returns it as a base64-encoded image.
    """
    labels = list(sections.keys())
    sizes = [section['score'] for section in sections.values()]
    
    # Map grades to colors for consistency with the report
    colors = []
    for section in sections.values():
        grade = section['grade'].lower()
        if grade == 'excellent':
            colors.append('#4CAF50')
        elif grade == 'good':
            colors.append('#2196F3')
        elif grade == 'average':
            colors.append('#FF9800')
        else:
            colors.append('#dc3545')

    plt.figure(figsize=(10, 10), facecolor='#121212')
    plt.pie(
        sizes,
        labels=labels,
        autopct='%1.1f%%',
        colors=colors,
        textprops={'color': "w", 'fontsize': 14}
    )
    plt.axis('equal')
    plt.tight_layout()

    buf = io.BytesIO()
    plt.savefig(buf, format='png', facecolor='#121212')
    buf.seek(0)
    encoded = base64.b64encode(buf.read()).decode('utf-8')
    buf.close()
    plt.close()
    return encoded


def generate_pie_chart_v2(sections):
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    import io, base64

    main_aspects = [
        "Format & Layout",
        "File Type & Parsing",
        "Section Headings & Structure",
        "Job-Title & Core Skills",
        "Dedicated Skills Section"
    ]

    labels = []
    sizes = []
    colors = ['#4CAF50', '#2196F3', '#FF9800', '#dc3545', '#673AB7']

    for aspect in main_aspects:
        if aspect in sections:
            score = sections[aspect].get('score', 0)
            labels.append(aspect)
            sizes.append(score)

    if not sizes or sum(sizes) == 0:
        return None

    fig, ax = plt.subplots(figsize=(10, 10), facecolor='#121212')
    wedges, texts, autotexts = ax.pie(
    sizes,
    autopct='%1.1f%%',
    colors=colors,
    textprops={'color': "white", 'fontsize': 20}
    )
    plt.axis('equal')

    plt.subplots_adjust(bottom=0.25)

    legend_labels = [f"{label}: {size}" for label, size in zip(labels, sizes)]
    ax.legend(
        wedges,
        legend_labels,
        title="Main Aspects",
        loc='lower center',
        bbox_to_anchor=(0.5, -0.5),
        fontsize=20,
        title_fontsize=20,
        frameon=False,
        labelcolor='white'
    )


    plt.tight_layout()

    buf = io.BytesIO()
    plt.savefig(buf, format='png', facecolor='#121212')
    buf.seek(0)
    encoded = base64.b64encode(buf.read()).decode('utf-8')
    buf.close()
    return encoded


# --- Resume ATS Scoring Functions ---
import io
from pdfminer.high_level import extract_text as pdf_extract_text
from docx import Document

def extract_resume_text(file) -> str:
    """
    Accepts InMemoryUploadedFile; returns plain text from PDF/DOCX/TXT.
    """
    name = (file.name or "").lower()
    data = file.read()
    file.seek(0)

    if name.endswith(".pdf"):
        try:
            return pdf_extract_text(io.BytesIO(data))
        except Exception:
            pass
    if name.endswith(".docx"):
        try:
            doc = Document(io.BytesIO(data))
            return "\n".join(p.text for p in doc.paragraphs)
        except Exception:
            pass
    # fallback: try decode as text
    try:
        return data.decode("utf-8", errors="ignore")
    except Exception:
        return ""

def normalize_text(t: str) -> str:
    return re.sub(r"\s+", " ", (t or "")).strip().lower()

def keyword_match_rate(text: str, target_keywords: list[str]) -> float:
    if not target_keywords:
        return 0.0
    t = normalize_text(text)
    hits = sum(1 for kw in target_keywords if kw.lower() in t)
    return hits / max(1, len(target_keywords))

# ===== Resume ATS (15 pts) =====
def ats_resume_scoring(metrics: dict) -> dict:
    """
    Resume (ATS Readiness) breakdown = 15 pts total, plus a normalized score out of 100.
    """
    b = {"items": []}
    total = 0
    MAX_ATS = 15

    # 1) Layout & structure — 3
    pts_layout = int(bool(metrics.get("sections_present"))) \
               + int(bool(metrics.get("single_column"))) \
               + int(bool(metrics.get("text_extractable")))
    b["items"].append({"name": "ATS-friendly layout & structure", "earned": pts_layout, "max": 3})
    total += pts_layout

    # 2) Action verbs & quantified results — 4
    av = float(metrics.get("action_verbs_per_bullet", 0.0))
    qr = float(metrics.get("quantified_bullets_ratio", 0.0))
    pts_actions = (2 if av >= 0.8 else 1 if av >= 0.5 else 0) \
                + (2 if qr >= 0.6 else 1 if qr >= 0.3 else 0)
    b["items"].append({"name": "Action verbs & quantified results", "earned": pts_actions, "max": 4})
    total += pts_actions

    # 3) Keyword alignment — 3
    kmr = float(metrics.get("keyword_match_rate", 0.0))
    pts_keywords = 3 if kmr >= 0.75 else 2 if kmr >= 0.5 else 1 if kmr >= 0.3 else 0
    b["items"].append({"name": "Job-relevant keyword alignment", "earned": pts_keywords, "max": 3})
    total += pts_keywords

    # 4) Brevity & conciseness — 2
    pages = int(metrics.get("pages", 2))
    avg_bullets = float(metrics.get("avg_bullets_per_job", 6.0))
    pts_brev = (1 if pages <= 2 else 0) + (1 if avg_bullets <= 7 else 0)
    b["items"].append({"name": "Brevity & conciseness", "earned": pts_brev, "max": 2})
    total += pts_brev

    # 5) Minimal jargon / repetition — 3
    rep = float(metrics.get("repetition_rate", 0.15))
    jar = float(metrics.get("jargon_rate", 0.2))
    usk = int(metrics.get("unique_skills_count", 8))
    pts_clean = (1 if rep <= 0.10 else 0) + (1 if jar <= 0.15 else 0) + (1 if usk >= 8 else 0)
    b["items"].append({"name": "Minimal jargon / repetition", "earned": pts_clean, "max": 3})
    total += pts_clean

    # Totals and normalized score
    b["subtotal"] = {"earned": total, "max": MAX_ATS}
    b["score_100"] = int(round((total / MAX_ATS) * 100))

    return b

# Role keyword lists (used for metrics + role match for non-tech)
ROLE_KEYWORDS = {
    # Technical
    "software engineer": ["python","java","javascript","react","node","docker","kubernetes","microservices","rest","graphql","aws","gcp","ci/cd","unit testing"],
    "data scientist": ["python","pandas","numpy","sklearn","tensorflow","pytorch","nlp","cv","statistics","sql","experiment","a/b testing","data visualization"],
    "devops engineer": ["ci/cd","docker","kubernetes","terraform","ansible","aws","gcp","azure","monitoring","prometheus","grafana","helm","sre"],
    "web developer": ["html","css","javascript","react","next.js","vue","node","express","rest","graphql","responsive","seo"],
    "mobile app developer": ["android","ios","kotlin","swift","flutter","react native","firebase","push notifications","play store","app store"],
    # Non-technical
    "human resources": ["recruitment","onboarding","payroll","employee engagement","hrms","policy","compliance","talent acquisition","grievance","training"],
    "marketing": ["seo","sem","campaign","content","email marketing","social media","analytics","branding","roi","conversion","google ads"],
    "sales": ["crm","pipeline","lead generation","negotiation","quota","prospecting","closing","upsell","cross-sell","demo"],
    "finance": ["budgeting","forecasting","reconciliation","audit","financial analysis","p&l","variance","sap","tally","excel"],
    "customer service": ["crm","zendesk","freshdesk","sla","csat","ticketing","call handling","escalation","knowledge base","communication"],
}

def derive_resume_metrics(resume_text: str, role_title: str) -> dict:
    t = normalize_text(resume_text)
    sections_present = any(k in t for k in ["experience","work history"]) and ("education" in t) and ("skills" in t)
    single_column = True
    text_extractable = len(t) > 0

    action_verbs = ["led","built","created","designed","implemented","developed","optimized","increased","reduced","launched","migrated","improved","delivered"]
    action_verb_hits = sum(len(re.findall(rf"(^|\n|•|\-)\s*({v})\b", resume_text, flags=re.I)) for v in action_verbs)
    bullets = max(1, len(re.findall(r"(\n•|\n-|\n\d+\.)", resume_text)))
    action_verbs_per_bullet = min(1.0, action_verb_hits / bullets)

    quantified_bullets_ratio = min(1.0, len(re.findall(r"\b\d+(\.\d+)?%?|\b(k|m|bn)\b", resume_text, flags=re.I)) / max(1, bullets))

    pages = max(1, round(len(resume_text) / 2000))
    avg_bullets_per_job = min(12.0, bullets / max(1, len(re.findall(r"\b(company|employer|experience)\b", t))))

    base_role = next((rk for rk in ROLE_KEYWORDS if rk in role_title.lower()), None)
    kws = ROLE_KEYWORDS.get(base_role, [])
    kmr = keyword_match_rate(resume_text, kws) if kws else 0.0

    repetition_rate = 0.08 if "responsible for" not in t else 0.18
    jargon_rate = 0.12 if "synergy" not in t and "leverage" not in t else 0.22

    unique_skills_count = len(set(re.findall(r"[a-zA-Z][a-zA-Z0-9\+\#\.\-]{1,20}", resume_text))) // 50
    unique_skills_count = max(0, min(unique_skills_count, 15))

    return {
        "sections_present": sections_present,
        "single_column": single_column,
        "text_extractable": text_extractable,
        "action_verbs_per_bullet": action_verbs_per_bullet,
        "quantified_bullets_ratio": quantified_bullets_ratio,
        "keyword_match_rate": kmr,
        "pages": pages,
        "avg_bullets_per_job": avg_bullets_per_job,
        "repetition_rate": repetition_rate,
        "jargon_rate": jargon_rate,
        "unique_skills_count": unique_skills_count,
    }

def calculate_screening_emphasis(github_score, linkedin_score, portfolio_score, resume_score, certifications_score):
    """
    Calculates the Screening Emphasis by Company Type based on the uploaded resume.
    
    Arguments:
        github_score: Score for GitHub presence
        linkedin_score: Score for LinkedIn optimization
        portfolio_score: Score for portfolio website
        resume_score: Score for ATS-friendly resume
        certifications_score: Score for certifications
        
    Returns:
        A dictionary with the final scores for each company type (MAANG, Startups, Mid-sized, Fortune 500).
    """
    
    # Company type weights based on the rubric
    company_weights = {
        "MAANG": {
            "GitHub": 22,
            "LinkedIn": 22,
            "Portfolio": 20,
            "Resume": 31,
            "Certifications": 5,
        },
        "Startups": {
            "GitHub": 30,
            "LinkedIn": 18,
            "Portfolio": 28,
            "Resume": 20,
            "Certifications": 4,
        },
        "Mid-sized": {
            "GitHub": 25,
            "LinkedIn": 22,
            "Portfolio": 23,
            "Resume": 24,
            "Certifications": 6,
        },
        "Fortune 500": {
            "GitHub": 18,
            "LinkedIn": 25,
            "Portfolio": 17,
            "Resume": 30,
            "Certifications": 10,
        }
    }

    # Calculate the total score for each company type
    scores_by_company_type = {}
    for company_type, weights in company_weights.items():
        total_score = 0
        total_score += github_score * (weights["GitHub"] / 100)
        total_score += linkedin_score * (weights["LinkedIn"] / 100)
        total_score += portfolio_score * (weights["Portfolio"] / 100)
        total_score += resume_score * (weights["Resume"] / 100)
        total_score += certifications_score * (weights["Certifications"] / 100)
        scores_by_company_type[company_type] = total_score
    
    return scores_by_company_type

import re

def extract_github_data(resume_text):
    """Extracts GitHub username and related details from the resume text."""
    github_url = re.search(r'https?://github\.com/([A-Za-z0-9_-]+)', resume_text)
    if github_url:
        return {"username": github_url.group(1)}
    return None
def extract_linkedin_data(resume_text):
    """Extracts LinkedIn username and related details from the resume text."""
    linkedin_url = re.search(r'https?://(www\.)?linkedin\.com/in/([A-Za-z0-9_-]+)', resume_text)
    if linkedin_url:
        return {"username": linkedin_url.group(2)}
    return None
def extract_portfolio_data(resume_text):
    """Extracts portfolio URL from the resume text."""
    portfolio_url = re.search(r'https?://[a-zA-Z0-9.-]+\.(me|tech|dev|xyz|site|vercel\.app|github\.io)', resume_text)
    if portfolio_url:
        return {"url": portfolio_url.group(0)}
    return None
def extract_certifications_data(resume_text):
    """Extracts certifications from the resume text."""
    certifications = []
    # Example: Extracting 'Certification Name, Issuer, Date'
    pattern = r'([A-Za-z0-9\s]+)\s*[-:]\s*(\w+)\s*[-:]\s*(\d{4})'
    matches = re.findall(pattern, resume_text)
    
    for match in matches:
        cert_name, issuer, year = match
        certifications.append({"name": cert_name.strip(), "issuer": issuer.strip(), "year": year.strip()})
    
    return certifications if certifications else None
def extract_resume_data(resume_text):
    """Extracts key sections (skills, experience, education, etc.) from the resume text."""
    resume_data = {}
    
    # Example: Extracting skills, experience, and education
    skills_pattern = r'(skills|technical skills|technologies)\s*[:\-]?\s*([A-Za-z, ]+)'
    experience_pattern = r'(experience|work history)\s*[:\-]?\s*([A-Za-z, ]+)'
    education_pattern = r'(education|degree)\s*[:\-]?\s*([A-Za-z, ]+)'

    skills = re.findall(skills_pattern, resume_text, flags=re.IGNORECASE)
    experience = re.findall(experience_pattern, resume_text, flags=re.IGNORECASE)
    education = re.findall(education_pattern, resume_text, flags=re.IGNORECASE)

    if skills:
        resume_data['skills'] = skills[0][1].strip()
    if experience:
        resume_data['experience'] = experience[0][1].strip()
    if education:
        resume_data['education'] = education[0][1].strip()

    return resume_data
def calculate_github_score(github_data):
    """Calculates GitHub score based on the presence of relevant repositories and activity."""
    if github_data:
        # Example: Scoring based on number of repos and activity
        repos = github_data.get("repos", [])
        score = 0
        for repo in repos:
            commits = repo.get("commits", 0)
            if commits > 5:
                score += 5
            else:
                score += 2
        return score
    return 0
def calculate_linkedin_score(linkedin_data):
    """Calculates LinkedIn score based on the profile information."""
    if linkedin_data:
        score = 0
        if 'headline' in linkedin_data and "role" in linkedin_data['headline'].lower():
            score += 5
        if 'experience' in linkedin_data:
            score += 10
        if 'education' in linkedin_data and "degree" in linkedin_data['education'].lower():
            score += 5
        return score
    return 0
def calculate_portfolio_score(portfolio_data):
    """Calculates Portfolio score based on the presence of relevant projects."""
    if portfolio_data:
        project_count = len(portfolio_data.get("projects", []))
        score = 0
        if project_count >= 3:
            score = 15
        elif project_count == 2:
            score = 10
        elif project_count == 1:
            score = 5
        return score
    return 0
def calculate_resume_score(resume_data):
    """Calculates Resume score based on ATS-friendliness, structure, and content."""
    score = 0
    if 'skills' in resume_data and resume_data['skills']:
        score += 5
    if 'experience' in resume_data and resume_data['experience']:
        score += 10
    if 'education' in resume_data and resume_data['education']:
        score += 5
    return score
def calculate_certifications_score(certifications_data):
    """Calculates Certifications score based on relevance and number of certifications."""
    if certifications_data:
        relevant_count = sum(1 for cert in certifications_data if cert.get("relevance_tag") == "relevant")
        return relevant_count * 5  # Example scoring, adjust based on actual requirements
    return 0


import re
from typing import Optional

_NAME_STOPWORDS = {
    "resume", "curriculum", "vitae", "cv", "profile", "contact", "contacts",
    "email", "mail", "phone", "mobile", "github", "linkedin", "portfolio",
    "experience", "summary", "about", "education", "skills", "projects",
    "work", "employment", "objective", "interests", "address", "location"
}

# Patterns we might see near the name line
_HEADER_NOISE_RE = re.compile(r"^\s*(resume|curriculum\s+vitae|cv)\s*$", re.I)
_EMAIL_RE = re.compile(r"[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}")
_PHONE_RE = re.compile(r"(\+?\d[\d\-\s().]{7,}\d)")
_LINK_RE = re.compile(r"(https?://|www\.)\S+", re.I)
_SECTION_RE = re.compile(r"^\s*(summary|about|profile|experience|education|skills|projects|work|employment|objective)\b", re.I)

# Name-like line: 2–4 tokens, mostly letters (allow diacritics & hyphens), capitalized words
def _looks_like_name_line(line: str) -> bool:
    if not line or len(line) > 80:
        return False
    if _HEADER_NOISE_RE.match(line):
        return False
    if _EMAIL_RE.search(line) or _PHONE_RE.search(line) or _LINK_RE.search(line):
        return False
    # remove commas and extra punctuation
    cleaned = re.sub(r"[^\w\s\-’'áéíóúàèìòùäëïöüñç]", " ", line, flags=re.I).strip()
    if not cleaned:
        return False
    tokens = [t for t in cleaned.split() if t]
    if not (1 < len(tokens) <= 4):
        return False
    # reject lines dominated by stopwords
    if sum(1 for t in tokens if t.lower() in _NAME_STOPWORDS) >= 1:
        return False
    # Heuristic: at least 2 tokens start with uppercase (or are ALLCAPS short)
    caplike = 0
    for t in tokens:
        if t[:1].isupper() or (t.isupper() and 2 <= len(t) <= 4):
            caplike += 1
    return caplike >= 2

def _normalize_name(s: str) -> str:
    s = re.sub(r"\s+", " ", s or "").strip(" -–—_.,\t\n\r")
    # Title case but keep internal caps in e.g. "McDonald", "O'Connor" approximately
    parts = []
    for p in s.split(" "):
        if not p:
            continue
        # preserve acronyms <= 4 chars
        if p.isupper() and len(p) <= 4:
            parts.append(p)
        else:
            parts.append(p[:1].upper() + p[1:].lower())
    # Remove leftover trailing role words occasionally attached
    while parts and parts[-1].lower() in {"developer","engineer","intern","student","analyst","manager"}:
        parts.pop()
    return " ".join(parts).strip()

def _name_from_email(resume_text: str) -> Optional[str]:
    m = _EMAIL_RE.search(resume_text or "")
    if not m:
        return None
    local = m.group(0).split("@", 1)[0]
    # common separators
    pieces = re.split(r"[._\-+]+", local)
    # filter out numbers/common words
    pieces = [p for p in pieces if p and not p.isdigit() and p.lower() not in {"mail","email","gmail","official","work","dev"}]
    if len(pieces) >= 2:
        guess = " ".join(pieces[:3])
        # avoid company-like tokens in the guess
        if not re.search(r"(inc|llc|ltd|corp|company|official)$", guess, re.I):
            return _normalize_name(guess)
    return None

def extract_applicant_name(resume_text: str) -> Optional[str]:
    """
    Attempts to extract the candidate's full name from resume text.
    Strategy:
      1) Scan the top ~15 non-empty lines for a name-like line (2–4 tokens, capitalized).
      2) Fall back to email local-part (e.g., "john.doe" -> "John Doe").
      3) Return None if confident extraction fails.
    """
    if not resume_text:
        return None

    lines = [ln.strip() for ln in (resume_text or "").splitlines()]
    # Ignore leading blank/noisy lines and hard section headers
    candidate_lines: list[str] = []
    for ln in lines:
        if not ln.strip():
            continue
        if _SECTION_RE.match(ln):
            # stop early if we hit a section before finding a name
            break
        candidate_lines.append(ln)
        if len(candidate_lines) >= 15:
            break

    # Check first, then second line blocks (some resumes start with big name, then contact row)
    for ln in candidate_lines[:6]:
        if _looks_like_name_line(ln):
            return _normalize_name(ln)

    # Sometimes the second or third line (after an all-caps name or logo) is the real name
    for ln in candidate_lines[6:15]:
        if _looks_like_name_line(ln):
            return _normalize_name(ln)

    # Fallback from email
    guess = _name_from_email(resume_text)
    if guess:
        return guess

    # Last resort: look for a capitalized two-token sequence near the very top line
    if candidate_lines:
        top = candidate_lines[0]
        m = re.search(r"\b([A-Z][a-zA-Z’'\-]{1,})\s+([A-Z][a-zA-Z’'\-]{1,})\b", top)
        if m:
            return _normalize_name(m.group(0))

    return None


import re
from typing import Optional

# regex for github profile links
_GITHUB_RE = re.compile(
    r"(?:https?://)?(?:www\.)?github\.com/([A-Za-z0-9](?:[A-Za-z0-9\-]{0,38}[A-Za-z0-9])?)",
    re.I
)

def extract_github_username(resume_text: str) -> Optional[str]:
    """
    Extracts a GitHub username from resume text.
    Rules:
      - Match valid GitHub profile URL forms.
      - Ignore trailing paths (/repo, /issues, etc.).
      - Username rules: 1–39 chars, alphanumeric + single hyphens (no leading/trailing hyphen).
    """
    if not resume_text:
        return None

    m = _GITHUB_RE.search(resume_text)
    if m:
        username = m.group(1)
        # normalize
        return username.strip()

    # fallback: if someone just wrote "github: johndoe"
    m2 = re.search(r"github[^a-z0-9]+([A-Za-z0-9][A-Za-z0-9\-]{0,38}[A-Za-z0-9])", resume_text, re.I)
    if m2:
        return m2.group(1).strip()

    return None



import re
from typing import Optional

# regex for leetcode profile urls
_LEETCODE_RE = re.compile(
    r"(?:https?://)?(?:www\.)?leetcode\.com/(u/)?([A-Za-z0-9_\-]+)",
    re.I
)

def extract_leetcode_username(resume_text: str) -> Optional[str]:
    """
    Extracts a LeetCode username from resume text.
    Rules:
      - Matches profile URLs (with or without '/u/').
      - Allows alphanumeric, underscore, and dash.
      - Returns just the username (not the whole URL).
    """
    if not resume_text:
        return None

    m = _LEETCODE_RE.search(resume_text)
    if m:
        username = m.group(2)
        return username.strip()

    # fallback: "LeetCode: johndoe123"
    m2 = re.search(r"leetcode[^a-z0-9]+([A-Za-z0-9_\-]{3,30})", resume_text, re.I)
    if m2:
        return m2.group(1).strip()

    return None

# utils.py (add these imports near your other utils imports)
import os
import re
from typing import List, Tuple
from urllib.parse import urlparse, urlunparse, parse_qsl, urlencode

try:
    from pypdf import PdfReader
except Exception:
    PdfReader = None  # will handle gracefully


# ---------- URL normalization helpers ----------

_URL_RE = re.compile(
    r'(?i)\b((?:https?://|www\.)[^\s<>()\[\]{}",;]+(?:\([^\s<>()]*\))*)'
)
_ALLOWED_SCHEMES = {"http", "https"}
_TRACKING_PARAMS = {
    "utm_source","utm_medium","utm_campaign","utm_term","utm_content","utm_id",
    "gclid","fbclid","ref","ref_src","igshid","mc_cid","mc_eid","mkt_tok","trk","trkCampaign","trkModule","yclid"
}
_TRAIL_STRIP = ".,);:!?'\"><}]}/"


def _clean_url(raw: str) -> str | None:
    if not raw:
        return None
    u = raw.strip().strip(_TRAIL_STRIP)
    if u.lower().startswith("www."):
        u = "https://" + u
    try:
        p = urlparse(u)
    except Exception:
        return None
    if p.scheme.lower() not in _ALLOWED_SCHEMES or not p.netloc:
        return None
    # strip trackers
    q = [(k, v) for k, v in parse_qsl(p.query, keep_blank_values=True) if k not in _TRACKING_PARAMS]
    cleaned = urlunparse((p.scheme, p.netloc, p.path or "", "", urlencode(q, doseq=True), ""))
    # trim trailing slash if path existed
    if cleaned.endswith("/") and p.path not in ("", "/"):
        cleaned = cleaned[:-1]
    return cleaned

def _dedupe_preserve_order(items: List[str]) -> List[str]:
    seen, out = set(), []
    for x in items:
        if x and x not in seen:
            seen.add(x); out.append(x)
    return out


# ---------- Plain text URL scanning ----------

def _extract_urls_from_text(text: str) -> List[str]:
    urls: List[str] = []
    for m in _URL_RE.finditer(text or ""):
        cleaned = _clean_url(m.group(0))
        if cleaned:
            urls.append(cleaned)
    return _dedupe_preserve_order(urls)


# ---------- PDF: clickable annotations + printed URLs ----------

def _extract_from_pdf(file_path: str) -> Tuple[List[str], str]:
    """
    Returns (urls, full_text) from a PDF:
      - Clickable link annotations (/Annots -> /A -> /URI)
      - Printed URLs in text layer
    """
    if PdfReader is None:
        # Library not available — best effort: no URLs, no text
        return [], ""

    urls: List[str] = []
    texts: List[str] = []
    try:
        reader = PdfReader(file_path)
    except Exception:
        return [], ""

    # 1) clickable annotations
    for page in reader.pages:
        try:
            if "/Annots" in page:
                for annot in page["/Annots"]:
                    obj = annot.get_object()
                    if obj.get("/Subtype") == "/Link":
                        a = obj.get("/A")
                        if a and a.get("/S") == "/URI":
                            uri = a.get("/URI")
                            cleaned = _clean_url(str(uri))
                            if cleaned:
                                urls.append(cleaned)
        except Exception:
            # ignore per-page annotation errors
            pass

    # 2) extract text & parse printed URLs
    try:
        for page in reader.pages:
            t = page.extract_text() or ""
            texts.append(t)
            # printed URLs on page
            urls.extend(_extract_urls_from_text(t))
    except Exception:
        # if text extraction fails, keep whatever we already got
        pass

    full_text = "\n".join(texts).strip()
    return _dedupe_preserve_order(urls), full_text


# ---------- TXT fallback ----------

def _extract_from_txt(file_path: str) -> Tuple[List[str], str]:
    try:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()
        return _extract_urls_from_text(text), text
    except Exception:
        return [], ""


# ---------- Public API ----------

def extract_links_combined(file_path: str) -> Tuple[List[str], str]:
    """
    Unified extractor used by your views for PDFs (primary), TXT (fallback),
    and best-effort for other file types (plain text URL scan if readable).

    Returns: (urls, full_text)
      - urls: list[str] of normalized, deduped URLs (http/https only)
      - full_text: the extracted text (PDF text layer or file content)
    """
    ext = (os.path.splitext(file_path)[1] or "").lower()

    if ext == ".pdf":
        return _extract_from_pdf(file_path)

    if ext in (".txt", ".md", ".csv", ".log"):
        return _extract_from_txt(file_path)

    # Best-effort generic fallback (attempt to read as text)
    try:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            txt = f.read()
        return _extract_urls_from_text(txt), txt
    except Exception:
        # Unknown binary types (e.g., docx) are intentionally handled elsewhere
        return [], ""

from typing import Optional
import docx

def extract_text_from_docx(path: str) -> Optional[str]:
    """
    Extract plain text from a .docx file.

    Args:
        path (str): Full path to the .docx file.

    Returns:
        str | None: All text joined with line breaks, or None if parsing fails.
    """
    try:
        doc = docx.Document(path)
        # Collect text from paragraphs
        text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])

        # Also collect text from tables (some resumes are table-heavy)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        text += "\n" + cell_text

        return text.strip() if text else None

    except Exception as e:
        print(f"[extract_text_from_docx] Failed to parse {path}: {e}")
        return None
    

import re
from typing import List, Tuple, Dict, Any
from urllib.parse import urlparse

# ---------------------------
# Helpers & constants (local)
# ---------------------------
_CERT_LINK_HOSTS = (
    "credly.com","youracclaim.com","accredible.com","badgr.com","openbadgepassport.com",
    "coursera.org","udemy.com","trailhead.salesforce.com","trailhead.salesforce",
    "aws.amazon.com","cloud.google.com","learn.microsoft.com","docs.microsoft.com",
    "oracle.com","education.oracle.com","datacamp.com","academy.databricks.com",
    "udacity.com","learn.udacity.com","tableau.com","certificates.tableau.com",
    "cisco.com","comptia.org","redhat.com","training.linuxfoundation.org","databricks.com",
    "snowflake.com","salesforce.com","trailhead.com"
)

_CERT_PROVIDERS = (
    "aws","amazon web services","azure","microsoft","google cloud","gcp",
    "coursera","udemy","udacity","datacamp","databricks","snowflake","tableau",
    "power bi","oracle","salesforce","cisco","red hat","linux foundation","ibm",
    "pmi","prince2","itil","scrum","comptia","kubernetes","ckad","cka","cks","sap",
    "okta","hashicorp","terraform","mongodb","redis","elastic"
)

_CERT_LEVEL_WORDS = (
    "associate","professional","expert","specialty","specialist","advanced",
    "foundational","practitioner","fundamentals","foundation","core"
)

# Common exam/credential patterns (keep simple, broad)
_EXAM_CODE_PAT = re.compile(
    r"\b(?:(?:AZ|DP|AI|PL|SC|MS|MB|MD|DA|SY|N|1Z0)-\d{2,4}|CKA|CKAD|CKS|SAA|DVA|SOA|SAP|MLS|DOP|CLF|DL|PDE)\b",
    re.I,
)
_CERT_KEYWORDS = re.compile(
    r"\b(certified|certification|certificate|credential|badge|license|licen[sc]e)\b",
    re.I
)
_BULLET_PREFIX = re.compile(r"^\s*(?:[-*•■●]|[0-9]+\.)\s*")

def _norm_text(s: str) -> str:
    return re.sub(r"\s{2,}", " ", (s or "").replace("\u00a0", " ")).strip(" -–—\t")

def _clean_cert_name(s: str) -> str:
    s = _BULLET_PREFIX.sub("", s or "")
    # Drop trailing IDs like "Credential ID: XYZ123"
    s = re.sub(r"\b(credential\s*id|id|license|licen[sc]e)\b.*$", "", s, flags=re.I)
    return _norm_text(s)[:180]

def _looks_like_cert_line(s: str) -> bool:
    if not s: return False
    t = s.lower()
    return bool(_CERT_KEYWORDS.search(s) or _EXAM_CODE_PAT.search(s) or any(p in t for p in _CERT_PROVIDERS))

def _domain_of(u: str) -> str:
    try:
        return (urlparse(u).netloc or "").lower()
    except Exception:
        return ""


# =========================================================
# 1) Count certifications (names only) from text + links
# =========================================================
import re
from typing import List, Dict, Tuple
from urllib.parse import urlparse

# Assumed helper functions, you can implement them as needed
import re
from typing import List, Dict, Tuple
from urllib.parse import urlparse

# Helper functions (implement as needed)
def _norm_text(text: str) -> str:
    """Normalize the text by stripping and lowering the case."""
    return text.strip().lower()

def _clean_cert_name(cert: str) -> str:
    """Clean up the certification name."""
    return cert.strip()

def _looks_like_cert_line(line: str) -> bool:
    """Detect if a line looks like a certification."""
    # Improved pattern for certifications
    return bool(re.search(r"(certification|certificate|certified|course|training|badge|diploma|credential)", line, re.IGNORECASE))

def count_only_certifications(resume_text: str) -> Tuple[int, List[Dict[str, str]]]:
    """
    Scans resume text for certifications and scores them dynamically based on their relevance.
    
    Args:
        resume_text: The text of the resume to scan for certifications.
        
    Returns:
        A tuple (count, certificates), where:
        - count is the number of unique certifications found
        - certificates is a list of dicts containing certification name and score.
    """
    certificates: List[Dict[str, str]] = []

    # --- From text (section + anywhere) ---
    lines = [ln.strip() for ln in (resume_text or "").splitlines()]
    in_cert_block = False
    for raw in lines:
        ln = _norm_text(raw)

        # Detect entering/leaving explicit section
        if re.match(r"^\s*(licenses?\s*&?\s*certifications?|certifications?|licenses?)\s*:?$", ln, re.I):
            in_cert_block = True
            continue
        if in_cert_block and (not ln or re.match(r"^(experience|education|projects?|skills?|profile|summary|achievements?)\s*:?\s*$", ln, re.I)):
            in_cert_block = False

        # Clean and validate the line
        cand = _clean_cert_name(ln)
        if (in_cert_block and cand) or _looks_like_cert_line(cand):
            # Only add reasonably short cert lines
            if 3 <= len(cand) <= 140:
                certificates.append({
                    "name": cand,
                    "source": "resume_text",
                    "score": 0  # Initial score to be determined
                })

    # Debugging - Print out the detected certifications
    print(f"Detected certifications: {[cert['name'] for cert in certificates]}")

    # --- Deduplicate (case/space-insensitive) preserving order ---
    seen = set()
    unique_certificates: List[Dict[str, str]] = []
    for cert in certificates:
        key = re.sub(r"[\s\-–—]+", " ", cert['name'].lower()).strip()
        if key and key not in seen:
            seen.add(key)
            unique_certificates.append(cert)

    # Debugging - Print out the unique certifications after deduplication
    print(f"Unique certifications after deduplication: {[cert['name'] for cert in unique_certificates]}")

    # --- Score the certificates ---
    for cert in unique_certificates:
        score = 0
        # Scoring based on the name of the certification (length and key phrases)
        if len(cert["name"]) > 50:
            score += 2  # Longer names might indicate more detailed, recognized certifications
        if "certification" in cert["name"].lower():
            score += 2  # Presence of the word 'certification' adds weight
        if "course" in cert["name"].lower():
            score += 1  # Presence of the word 'course' adds a moderate weight
        if "badge" in cert["name"].lower():
            score += 1  # Presence of 'badge' adds weight as it could be a recognized platform

        cert["score"] = score

    # Final debugging output for scoring
    print(f"Scored certifications: {[cert['name'] + ' (score: ' + str(cert['score']) + ')' for cert in unique_certificates]}")

    return len(unique_certificates), unique_certificates

# 

# =========================================================
# 2) Suggest role-aware certifications (strings only)
# =========================================================
def suggest_role_certifications(
    role_text: str,
    job_description: str = "",
    resume_text: str = "",
    existing_cert_lines: List[str] | None = None,
    max_items: int = 6,
) -> List[str]:
    """
    Suggest reputable certs aligned to role/JD. Avoids suggesting ones already present.

    Returns up to `max_items` suggestion strings (no links).
    """
    role_blob = " ".join([role_text or "", job_description or "", resume_text or ""]).lower()

    # Buckets with priority ordering inside each
    SUGGEST = {
        "cloud": [
            "AWS Certified Cloud Practitioner (CLF-C02)",
            "AWS Certified Solutions Architect – Associate (SAA-C03)",
            "Microsoft Certified: Azure Fundamentals (AZ-900)",
            "Microsoft Certified: Azure Administrator Associate (AZ-104)",
            "Google Cloud Digital Leader",
            "Google Associate Cloud Engineer",
        ],
        "devops": [
            "Certified Kubernetes Administrator (CKA)",
            "Certified Kubernetes Application Developer (CKAD)",
            "HashiCorp Terraform Associate (003)",
            "AWS Certified Developer – Associate (DVA-C02)",
            "Microsoft DevOps Engineer Expert (AZ-400)",
        ],
        "data_eng": [
            "Google Professional Data Engineer",
            "Azure Data Engineer Associate (DP-203)",
            "Databricks Data Engineer Associate",
            "SnowPro Core Certification",
            "AWS Certified Data Analytics – Specialty (DAS-C01)",
        ],
        "data_sci_ml": [
            "Azure AI Fundamentals (AI-900)",
            "Azure Data Scientist Associate (DP-100)",
            "Google Professional Machine Learning Engineer",
            "AWS Certified Machine Learning – Specialty (MLS-C01)",
            "TensorFlow Developer Certificate",
        ],
        "security": [
            "CompTIA Security+ (SY0-701)",
            "AWS Certified Security – Specialty (SCS-C02)",
            "Azure Security Engineer Associate (AZ-500)",
            "Google Professional Cloud Security Engineer",
        ],
        "frontend": [
            "Meta Front-End Developer (Professional Certificate, Coursera)",
            "AWS Certified Cloud Practitioner (CLF-C02)",
        ],
        "backend": [
            "Oracle Certified Professional: Java SE 11 Developer (1Z0-819)",
            "AWS Certified Developer – Associate (DVA-C02)",
            "Docker Certified Associate (DCA)",
        ],
        "qa_sdet": [
            "ISTQB Certified Tester – Foundation Level (CTFL)",
            "Certified Kubernetes Application Developer (CKAD)",
            "AWS Certified Developer – Associate (DVA-C02)",
        ],
        "product": [
            "Professional Scrum Product Owner I (PSPO I)",
            "Certified Scrum Product Owner (CSPO)",
            "Pragmatic Institute PMC Level I",
        ],
        "analytics_bi": [
            "Microsoft Power BI Data Analyst Associate (PL-300)",
            "Tableau Desktop Specialist",
            "Google Business Intelligence Professional Certificate",
        ],
        "generic": [
            "AWS Certified Cloud Practitioner (CLF-C02)",
            "Microsoft Certified: Azure Fundamentals (AZ-900)",
            "Google Cloud Digital Leader",
        ],
    }

    # Simple keyword routing
    bucket_scores = {
        "cloud": int(any(k in role_blob for k in ["cloud", "aws", "azure", "gcp", "solutions architect"])),
        "devops": int(any(k in role_blob for k in ["devops", "sre", "platform", "kubernetes", "docker", "terraform"])),
        "data_eng": int(any(k in role_blob for k in ["data engineer", "etl", "elt", "spark", "airflow", "dbt", "snowflake", "bigquery", "redshift"])),
        "data_sci_ml": int(any(k in role_blob for k in ["data scientist", "machine learning", "ml engineer", "mlops", "ai " , "deep learning"])),
        "security": int(any(k in role_blob for k in ["security", "infosec", "cloud security", "zero trust", "blue team", "red team"])),
        "frontend": int(any(k in role_blob for k in ["frontend", "front-end", "react", "javascript", "ui engineer"])),
        "backend": int(any(k in role_blob for k in ["backend", "back-end", "java", "spring", "node.js", "python developer"])),
        "qa_sdet": int(any(k in role_blob for k in ["qa", "quality", "sdet", "test automation", "tester"])),
        "product": int(any(k in role_blob for k in ["product manager", "pm ", "product owner"])),
        "analytics_bi": int(any(k in role_blob for k in ["bi", "business intelligence", "power bi", "tableau", "data analyst"])),
    }

    # Ordered buckets by relevance, then some generic cloud as catch-all
    ordered_buckets = [k for k, v in sorted(bucket_scores.items(), key=lambda kv: kv[1], reverse=True) if v] or ["generic"]

    # Flatten suggestions following buckets order
    flat: List[str] = []
    for b in ordered_buckets:
        for c in SUGGEST[b]:
            if c not in flat:
                flat.append(c)
    # Add some generic ones if still room
    if "generic" not in ordered_buckets:
        for c in SUGGEST["generic"]:
            if c not in flat:
                flat.append(c)

    # Remove already-owned certs (fuzzy)
    existing = existing_cert_lines or []
    def _key(s: str) -> str:
        return re.sub(r"[^a-z0-9]+", " ", (s or "").lower()).strip()

    owned_keys = {_key(x) for x in existing}
    out: List[str] = []
    for cert in flat:
        if _key(cert) not in owned_keys:
            out.append(cert)
        if len(out) >= max_items:
            break

    return out


# =========================================================
# 3) Score LinkedIn PUBLIC HTML (0–18)
# =========================================================
def score_linkedin_public_html(html: str, url: str = "", resume_text: str = "") -> Tuple[int, List[str], List[str]]:
    """
    Heuristically score a PUBLIC LinkedIn profile page's HTML (0–18).
    This does *not* log in—only parses the HTML string you already fetched.

    Signals (max 18 pts):
      +5  Experience section present
      +3  Education section present
      +3  Skills section present
      +0/1/3  Endorsements mentioned (>=1 → +1, >=3 → +3)
      +2  Recommendations present
      +0/1/2  Followers / connections (>=100 → +1, >=300 → +2)

    Returns:
      (score: int, rationales: List[str], evidence_links: List[str])
    """
    if not html:
        return 0, ["LinkedIn profile HTML not available (likely behind login)."], [url] if url else []

    rats: List[str] = []
    ev = [url] if url else []
    low = html.lower()

    has_exp = any(re.search(p, low) for p in (
        r">experience<", r'\bexperience\b.{0,40}section', r'id="experience"', r'data-test-?section[^>]*experience'
    ))
    has_edu = any(re.search(p, low) for p in (
        r">education<", r'\beducation\b.{0,40}section', r'id="education"', r'data-test-?section[^>]*education'
    ))
    has_skills = any(re.search(p, low) for p in (
        r">skills<", r'\bskills\b.{0,40}section', r'id="skills"', r'data-test-?section[^>]*skills'
    ))
    has_recs = bool(re.search(r"\brecommendation[s]?\b", low))
    endorsements_mentions = len(re.findall(r"\bendorsement[s]?\b", low))

    followers_num = 0
    m_follow = re.search(r"(\d[\d,]{1,6})\s+(?:followers|connections)", low)
    if m_follow:
        try:
            followers_num = int(m_follow.group(1).replace(",", ""))
        except Exception:
            followers_num = 0

    score = 0
    if has_exp: score += 5; rats.append("Experience section present (+5).")
    else:       rats.append("No visible Experience section (0).")

    if has_edu: score += 3; rats.append("Education section present (+3).")
    else:       rats.append("No visible Education section (0).")

    if has_skills: score += 3; rats.append("Skills section present (+3).")
    else:          rats.append("No visible Skills section (0).")

    if endorsements_mentions >= 3:
        score += 3; rats.append(f"Endorsements mentioned ≈ {endorsements_mentions} (+3).")
    elif endorsements_mentions >= 1:
        score += 1; rats.append(f"Endorsements mentioned ≈ {endorsements_mentions} (+1).")
    else:
        rats.append("No endorsements mentions (0).")

    if has_recs: score += 2; rats.append("Recommendations present (+2).")
    else:        rats.append("No recommendations mentions (0).")

    if followers_num >= 300:
        score += 2; rats.append(f"Connections/Followers ≈ {followers_num} (+2).")
    elif followers_num >= 100:
        score += 1; rats.append(f"Connections/Followers ≈ {followers_num} (+1).")
    else:
        rats.append("Connections/Followers not detected or <100 (0).")

    score = min(18, max(0, score))
    return score, rats, ev


# profile_scoring.py
# ------------------------------------------------------------
# GitHub scoring (API) -> 27 pts
# LinkedIn scoring (Gemini) -> 18 pts (7 subscores)
# LeetCode (Gemini) -> auxiliary 0–9
# NEW: Mandatory resume link extraction (PDF/DOCX/TXT) + metadata fetching
# ------------------------------------------------------------

from __future__ import annotations
import os, re, json, math, time, requests
from typing import Dict, Any, List, Tuple, Optional
from urllib.parse import urlparse, urlunparse, parse_qsl, urlencode

# ========== Optional .env ==========
try:
    from dotenv import load_dotenv
    load_dotenv()
except Exception:
    pass

# ========== Gemini ==========
import google.generativeai as genai
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "")
if GEMINI_API_KEY:
    genai.configure(api_key=GEMINI_API_KEY)
_GEM_MODEL_NAME = os.getenv("GEMINI_MODEL", "gemini-1.5-flash")

# ========== Common utils ==========
def _norm(s: Optional[str]) -> str:
    return (s or "").strip()

def _dedupe(seq: List[str], limit: int = 8) -> List[str]:
    seen, out = set(), []
    for x in seq:
        if x and x not in seen:
            seen.add(x); out.append(x)
        if len(out) >= limit: break
    return out

def _domain(url: str) -> str:
    try:
        return urlparse(url).netloc.lower()
    except Exception:
        return ""

# ============================================================
# 1) GitHub — 27 pts (via API)
# Bands:
#   2–4: link present only
#   7–10: 1 repo, minimal activity
#   13–18: 1 repo significant contributions
#   20–27: ≥3 active/hosted projects, varied skills, docs, tests, demos
# ============================================================

GH_API = "https://api.github.com"

def _gh_headers(token: Optional[str]) -> Dict[str, str]:
    headers = {
        "Accept": "application/vnd.github+json",
        "X-GitHub-Api-Version": "2022-11-28",
        "User-Agent": "applywizz-profile-scoring"
    }
    if token:
        headers["Authorization"] = f"Bearer {token}"
    return headers

def _gh_get(url: str, token: Optional[str], params: Dict[str, Any] = None) -> requests.Response:
    """
    GET with PAT if provided; if we get a 401, retry once unauthenticated.
    """
    params = params or {}
    # First attempt (auth if token present)
    r = requests.get(url, headers=_gh_headers(token), params=params, timeout=30)
    if r.status_code == 401 and token:
        # Retry unauthenticated to avoid hard-crash; rate limits will be lower
        r = requests.get(url, headers=_gh_headers(None), params=params, timeout=30)
    r.raise_for_status()
    return r

def _estimate_commit_count(owner: str, repo: str, token: str) -> int:
    """
    Efficient commit count: request per_page=1 then read 'last' page from Link header.
    """
    url = f"{GH_API}/repos/{owner}/{repo}/commits"
    r = _gh_get(url, token, params={"per_page": 1})
    link = r.headers.get("Link", "")
    m = re.search(r'&page=(\d+)>; rel="last"', link)
    if m:
        return int(m.group(1))
    data = r.json()
    return 1 if isinstance(data, list) and data else 0

def _has_readme(owner: str, repo: str, token: str) -> bool:
    url = f"{GH_API}/repos/{owner}/{repo}/readme"
    r = requests.get(url, headers=_gh_headers(token), timeout=30)
    return r.status_code == 200

def _has_tests(owner: str, repo: str, token: str) -> bool:
    common = ["tests", "test", "spec", "src/test", "cypress", "__tests__"]
    headers = _gh_headers(token)
    for path in common:
        url = f"{GH_API}/repos/{owner}/{repo}/contents/{path}"
        rr = requests.get(url, headers=headers, timeout=20)
        if rr.status_code == 200:
            return True
        if rr.status_code in (403, 429):
            time.sleep(0.5)
    return False

def _get_languages(owner: str, repo: str, token: str) -> List[str]:
    url = f"{GH_API}/repos/{owner}/{repo}/languages"
    try:
        r = _gh_get(url, token)
        data = r.json() or {}
        return list(data.keys())
    except Exception:
        return []

def _repo_demo_url(repo_obj: Dict[str, Any]) -> str:
    if repo_obj.get("has_pages") and _norm(repo_obj.get("homepage")):
        return repo_obj["homepage"]
    return _norm(repo_obj.get("homepage"))

def fetch_github_profile(username: str, token: str) -> Dict[str, Any]:
    u = _gh_get(f"{GH_API}/users/{username}", token).json()
    repos = []
    page = 1
    while True:
        rr = _gh_get(f"{GH_API}/users/{username}/repos", token, params={"per_page": 100, "page": page, "sort": "updated"})
        chunk = rr.json()
        if not chunk: break
        repos.extend(chunk)
        page += 1
        if page > 10: break  # safety
    return {"user": u, "repos": repos}

def score_github_via_api(username: str, token: str) -> Tuple[int, List[str], List[str], Dict[str, Any]]:
    payload = fetch_github_profile(username, token)
    repos_raw = payload.get("repos", []) or []
    profile_url = payload.get("user", {}).get("html_url", f"https://github.com/{username}")
    evidence = [profile_url]

    enriched = []
    for r in repos_raw[:15]:
        if r.get("fork"):
            continue
        owner = r["owner"]["login"]
        name = r["name"]

        try:
            commits = _estimate_commit_count(owner, name, token)
        except Exception:
            commits = r.get("forks_count", 0)

        try:
            readme = _has_readme(owner, name, token)
        except Exception:
            readme = False

        try:
            tests = _has_tests(owner, name, token)
        except Exception:
            tests = False

        langs = _get_languages(owner, name, token)
        demo = _repo_demo_url(r)

        enriched.append({
            "name": name,
            "html_url": r.get("html_url"),
            "stars": int(r.get("stargazers_count", 0)),
            "commits": commits,
            "has_readme": readme,
            "has_tests": tests,
            "languages": langs,
            "demo_url": demo,
            "archived": r.get("archived", False),
            "has_pages": r.get("has_pages", False),
            "pushed_at": r.get("pushed_at"),
        })

    active = [x for x in enriched if not x["archived"] and x["commits"] >= 6]
    significant = [x for x in active if x["commits"] >= 30 and (x["has_readme"] or x["has_tests"])]
    hosted = [x for x in active if _norm(x["demo_url"])]
    langs_div = set(l.lower() for x in active for l in x["languages"])

    if not enriched:
        return 0, ["No repositories found."], evidence, {"repos": enriched}

    if len(active) == 0:
        score = 3
        r = ["Profile present but no active original repos."]
    elif len(active) == 1:
        a = active[0]
        if a["commits"] >= 30 and (a["has_readme"] or a["has_tests"]):
            score = 15  # 13–18 band midpoint
            r = ["One repo with meaningful commits and docs/tests."]
        else:
            score = 8  # 7–10 band
            r = ["One active repo with minimal depth."]
    else:
        base = 20
        bumps = 0
        if len(significant) >= 2: bumps += 2
        if len(hosted) >= 1:      bumps += 2
        if len(langs_div) >= 3:   bumps += 2
        score = min(27, base + bumps)
        r = []
        if len(significant) >= 2: r.append("Multiple repos with significant commits/docs.")
        if len(hosted) >= 1:      r.append("Includes live demos/hosted projects.")
        if len(langs_div) >= 3:   r.append("Diverse stacks/languages.")

    for x in hosted[:3]:
        evidence.append(x["demo_url"])
    for x in active[:2]:
        evidence.append(x["html_url"])
    evidence = _dedupe(evidence)

    details = {
        "profile_url": profile_url,
        "active_count": len(active),
        "significant_count": len(significant),
        "hosted_count": len(hosted),
        "language_diversity": len(langs_div),
        "repos_enriched": enriched,
    }
    return score, (r or ["GitHub evaluated."]), evidence, details

# ============================================================
# 2) LinkedIn — 18 pts via Gemini (7 subscores)
# ============================================================

_LINKEDIN_SYSTEM = """You are a strict evaluator for LinkedIn profiles.
Score each sub-dimension per rubric and return only JSON. Cap totals exactly as specified.
Sub-dimensions:
- headline (0-2): role + relevant keywords
- about (0-3): narrative w/ achievements & goals
- experience (0-4): quantified results; consistent dates
- projects (0-3): featured/pinned with links
- education (0-2): complete & relevant
- skills (0-2): relevant & endorsed
- certificates (0-2): relevant and visible

Rules:
- Do not invent data; award conservatively when unclear.
- Never exceed category maxima or subscore maxima.
- Keep rationales terse (≤20 words)."""

def score_linkedin_with_gemini(linkedin_json: Dict[str, Any]) -> Tuple[int, Dict[str, int], List[str], List[str], Dict[str, Any]]:
    if not GEMINI_API_KEY:
        raise RuntimeError("GEMINI_API_KEY not set. Set it to enable LinkedIn scoring.")

    user_msg = json.dumps({"linkedin": linkedin_json}, ensure_ascii=False)
    model = genai.GenerativeModel(_GEM_MODEL_NAME)
    prompt = f"""{_LINKEDIN_SYSTEM}

Return strictly this JSON:
{{
  "subscores": {{
    "headline": int, "about": int, "experience": int, "projects": int,
    "education": int, "skills": int, "certificates": int
  }},
  "rationales": [string],   // 1–3 bullets
  "evidence": [string]      // links present inside linkedin_json if any
}}
"""
    resp = model.generate_content([{"role":"user","parts":[user_msg + "\n\n" + prompt]}])
    text = resp.text or "{}"
    try:
        obj = json.loads(re.search(r"\{[\s\S]+\}$", text).group(0))
    except Exception:
        obj = {"subscores": {"headline":0,"about":0,"experience":0,"projects":0,"education":0,"skills":0,"certificates":0},
               "rationales": ["Gemini output parsing failed."], "evidence": []}

    # cap per subscore maxima
    caps = {"headline":2,"about":3,"experience":4,"projects":3,"education":2,"skills":2,"certificates":2}
    subs_raw = obj.get("subscores", {}) or {}
    subs = {k:int(max(0, min(int(subs_raw.get(k,0)), caps[k]))) for k in caps}
    total = min(18, sum(subs.values()))
    rats = obj.get("rationales") or ["LinkedIn evaluated."]
    ev = _dedupe(obj.get("evidence") or [], limit=6)
    details = {"raw": obj}
    return total, subs, rats[:3], ev, details

# ============================================================
# 3) LeetCode — Auxiliary via Gemini (0–9)
# ============================================================

LEETCODE_GQL = "https://leetcode.com/graphql"

def fetch_leetcode_stats(username: str) -> Dict[str, Any]:
    q = """
    query userProfile($username: String!) {
      allQuestionsCount { difficulty count }
      matchedUser(username: $username) {
        username
        submitStats: submitStatsGlobal { acSubmissionNum { difficulty count submissions } }
        profile { ranking reputation }
        problemsSolvedBeatsStats { difficulty percentage }
        contestBadge { name }
        userCalendar { streak }
      }
    }
    """
    r = requests.post(LEETCODE_GQL, json={"query": q, "variables": {"username": username}}, timeout=30)
    r.raise_for_status()
    return r.json()

_LEETCODE_SYSTEM = """You are scoring LeetCode proficiency. Output only JSON.
Inputs include: total solved by difficulty, acceptance counts, streak, ranking, reputation, contest badge, and beats %
Map to a 0–9 score:
- 0–1: little/no activity
- 2–4: some easy/medium solves
- 5–7: consistent medium/hard; streaks; decent ranking
- 8–9: high volume; contest performance; strong ranking
Return rationale (≤20 words)."""

def score_leetcode_with_gemini(stats_or_username: str | Dict[str, Any]) -> Tuple[int, List[str], Dict[str, Any]]:
    if not GEMINI_API_KEY:
        raise RuntimeError("GEMINI_API_KEY not set. Set it to enable LeetCode scoring.")
    if isinstance(stats_or_username, str):
        raw = fetch_leetcode_stats(stats_or_username)
    else:
        raw = stats_or_username

    model = genai.GenerativeModel(_GEM_MODEL_NAME)
    prompt = f"""{_LEETCODE_SYSTEM}

Return strictly this JSON:
{{
  "score": int,             // 0-9
  "rationales": [string]    // 1–2 bullets
}}
"""
    resp = model.generate_content([{"role":"user","parts":[json.dumps(raw, ensure_ascii=False) + "\n\n" + prompt]}])
    text = resp.text or "{}"
    try:
        obj = json.loads(re.search(r"\{[\s\S]+\}$", text).group(0))
    except Exception:
        obj = {"score": 0, "rationales": ["Gemini output parsing failed."]}

    score = int(max(0, min(obj.get("score", 0), 9)))
    rats = obj.get("rationales") or ["LeetCode evaluated."]
    details = {"raw": raw, "gemini": obj}
    return score, rats[:2], details

# ============================================================
# 4) NEW — Resume Link Extraction + Metadata
# - Extract clickable anchors from PDF annotations & DOCX hyperlink rels
# - Also parse plain text URLs in PDF/DOCX/TXT
# - Normalize and dedupe
# - Fetch "what the link is about" = title + description (+ type)
# ============================================================

# ---- Dependencies for file parsing ----
from pypdf import PdfReader
from docx import Document
from bs4 import BeautifulSoup

# URL regex & normalization
URL_RE = re.compile(r'(?i)\b((?:https?://|www\.)[^\s<>()\[\]{}",;]+(?:\([^\s<>()]*\))*)')
TRAIL_STRIP = '.,);:!?\'"<>]}'
ALLOWED_SCHEMES = {"http", "https"}
TRACKING_PARAMS = {
    "utm_source","utm_medium","utm_campaign","utm_term","utm_content","utm_id",
    "gclid","fbclid","ref","ref_src","igshid","mc_cid","mc_eid","mkt_tok","trk","trkCampaign","trkModule","yclid"
}

def _clean_url(u: str) -> str | None:
    if not u: return None
    u = u.strip().strip(TRAIL_STRIP)
    if u.lower().startswith("www."):
        u = "https://" + u
    try:
        p = urlparse(u)
    except Exception:
        return None
    if p.scheme.lower() not in ALLOWED_SCHEMES: return None
    if not p.netloc: return None
    q = [(k,v) for (k,v) in parse_qsl(p.query, keep_blank_values=True) if k not in TRACKING_PARAMS]
    cleaned = urlunparse((p.scheme, p.netloc, p.path or "", "", urlencode(q, doseq=True), ""))
    if cleaned.endswith("/") and p.path not in ("","/"):
        cleaned = cleaned[:-1]
    return cleaned

def _extract_urls_from_text(text: str) -> List[str]:
    urls = []
    for m in URL_RE.finditer(text or ""):
        cleaned = _clean_url(m.group(0))
        if cleaned: urls.append(cleaned)
    return _dedupe(urls, limit=None)

def _extract_from_pdf(file_path: str) -> List[str]:
    urls: List[str] = []
    reader = PdfReader(file_path)
    # 1) real clickable link annotations
    for page in reader.pages:
        try:
            if "/Annots" in page:
                for annot in page["/Annots"]:
                    obj = annot.get_object()
                    if obj.get("/Subtype") == "/Link":
                        a = obj.get("/A")
                        if a and a.get("/S") == "/URI":
                            uri = a.get("/URI")
                            cleaned = _clean_url(str(uri))
                            if cleaned: urls.append(cleaned)
        except Exception:
            pass
    # 2) textual URLs
    try:
        for page in reader.pages:
            txt = page.extract_text() or ""
            urls.extend(_extract_urls_from_text(txt))
    except Exception:
        pass
    return _dedupe(urls, limit=None)

def _extract_from_docx(file_path: str) -> List[str]:
    urls: List[str] = []
    doc = Document(file_path)
    # 1) hyperlink relationships
    try:
        for rel in doc.part.rels.values():
            if rel.reltype.endswith("/hyperlink"):
                cleaned = _clean_url(rel.target_ref)
                if cleaned: urls.append(cleaned)
    except Exception:
        pass
    # 2) text in paragraphs & tables
    def _scan(container):
        for p in container.paragraphs:
            urls.extend(_extract_urls_from_text(p.text or ""))
        for t in container.tables:
            for row in t.rows:
                for cell in row.cells:
                    _scan(cell)
    _scan(doc)
    return _dedupe(urls, limit=None)

def _extract_from_txt(file_path: str) -> List[str]:
    try:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return _extract_urls_from_text(f.read())
    except Exception:
        return []

def extract_resume_links(file_path: str, mime_type: str | None = None, limit: int | None = 50) -> List[str]:
    ext = (os.path.splitext(file_path)[1] or "").lower()
    urls: List[str] = []
    try:
        if (mime_type and "pdf" in mime_type) or ext == ".pdf":
            urls = _extract_from_pdf(file_path)
        elif (mime_type and "word" in mime_type) or mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" or ext == ".docx":
            urls = _extract_from_docx(file_path)
        elif ext in (".txt", ".md"):
            urls = _extract_from_txt(file_path)
        else:
            # best-effort: try text parse
            urls = _extract_from_txt(file_path)
    except Exception:
        pass
    urls = _dedupe(urls, limit=limit or 50)
    return urls

# ---- Link metadata (what is the link about?) ----

DEFAULT_UA = "Mozilla/5.0 (compatible; ApplyWizz-Profile-Scorer/1.0; +https://example.com)"
HTML_TIMEOUT = 12

def _fetch_html(url: str) -> Optional[str]:
    try:
        r = requests.get(url, headers={"User-Agent": DEFAULT_UA}, timeout=HTML_TIMEOUT, allow_redirects=True)
        ct = (r.headers.get("Content-Type") or "").lower()
        if "text/html" not in ct and "application/xhtml" not in ct:
            return None
        if r.status_code >= 400:
            return None
        return r.text
    except Exception:
        return None

def _extract_title_desc(html: str) -> Tuple[str, str]:
    soup = BeautifulSoup(html, "lxml")
    # Try OG/Twitter meta first
    title = (
        (soup.find("meta", property="og:title") or {}).get("content") or
        (soup.find("meta", attrs={"name":"twitter:title"}) or {}).get("content") or
        (soup.title.string if soup.title else "")
    )
    desc = (
        (soup.find("meta", property="og:description") or {}).get("content") or
        (soup.find("meta", attrs={"name":"description"}) or {}).get("content") or
        (soup.find("meta", attrs={"name":"twitter:description"}) or {}).get("content") or
        ""
    )
    return (_norm(title), _norm(desc))

def classify_link_type(url: str) -> str:
    d = _domain(url)
    if "github.com" in d:       return "github"
    if "linkedin.com" in d:     return "linkedin"
    if "leetcode.com" in d:     return "leetcode"
    if "x.com" in d or "twitter.com" in d: return "social"
    if "medium.com" in d or "dev.to" in d or "hashnode" in d: return "blog"
    if "youtube.com" in d or "youtu.be" in d: return "video"
    if "kaggle.com" in d:       return "kaggle"
    if "notion.site" in d or "notion.so" in d: return "notion"
    if "readthedocs" in d or d.endswith(".readthedocs.io"): return "docs"
    return "other"

def describe_links(urls: List[str], with_fetch: bool = True, limit: int = 50) -> List[Dict[str, Any]]:
    """
    Returns: [{url, domain, type, title, description}]
    If with_fetch is True, fetch HTML and extract title+description for each link (best-effort).
    """
    out: List[Dict[str, Any]] = []
    for u in urls[:limit]:
        item = {
            "url": u,
            "domain": _domain(u),
            "type": classify_link_type(u),
            "title": "",
            "description": ""
        }
        if with_fetch:
            html = _fetch_html(u)
            if html:
                title, desc = _extract_title_desc(html)
                item["title"] = title
                item["description"] = desc
        out.append(item)
    return out

def process_resume_links(file_path: str, mime_type: Optional[str] = None, fetch_metadata: bool = True, limit: int = 50) -> List[Dict[str, Any]]:
    """
    High-level API:
      1) Extract all clickable/anchor links from resume (PDF/DOCX/TXT)
      2) Normalize & dedupe
      3) Optionally fetch 'what the link is about' (title/description)
    Returns list of {url, domain, type, title, description}
    """
    urls = extract_resume_links(file_path, mime_type=mime_type, limit=limit)
    return describe_links(urls, with_fetch=fetch_metadata, limit=limit)

# ============================================================
# Example usage (you can remove this in production)
# ============================================================

def example_run():
    # ---- GitHub ----
    gh_token = os.getenv("GITHUB_TOKEN")
    gh_user = "torvalds"  # replace with candidate's username
    if gh_token:
        gh_score, gh_r, gh_ev, gh_details = score_github_via_api(gh_user, gh_token)
        print("GitHub score:", gh_score, "/27")
        print("Rationales:", gh_r)
        print("Evidence:", gh_ev)
    else:
        print("Set GITHUB_TOKEN to test GitHub scoring.")

    # ---- LinkedIn (Gemini) ----
    linkedin_json = {
        "headline": "Data Engineer | Python • SQL • Spark • AWS",
        "about": "Built pipelines with measurable impact; seeking to scale platforms.",
        "experience": [{"title":"DE","desc":"Reduced runtime 40%; processed 120M rows","dates":"2022–2025"}],
        "projects": [{"title":"ETL Modernization","url":"https://example.com"}],
        "education": [{"degree":"MS"}],
        "skills": [{"name":"Python","endorsements":12},{"name":"SQL","endorsements":5}],
        "certificates": [{"name":"AWS SA"}]
    }
    if GEMINI_API_KEY:
        li_total, li_subs, li_r, li_ev, li_details = score_linkedin_with_gemini(linkedin_json)
        print("LinkedIn score:", li_total, "/18", "subs:", li_subs)
        print("Rationales:", li_r)
    else:
        print("Set GEMINI_API_KEY to enable LinkedIn scoring.")

    # ---- LeetCode (Gemini) ----
    if GEMINI_API_KEY:
        lc_score, lc_r, lc_details = score_leetcode_with_gemini("leetcode")  # replace with username
        print("LeetCode score (aux):", lc_score, "/9")
        print("Rationales:", lc_r)

    # ---- NEW: Resume links + metadata ----
    # Example: file_path = "/path/to/resume.pdf"
    file_path = os.getenv("RESUME_PATH")  # set for demo
    if file_path and os.path.exists(file_path):
        items = process_resume_links(file_path, mime_type=None, fetch_metadata=True, limit=30)
        print("Extracted resume links:")
        for it in items:
            print(f"- {it['url']} [{it['type']}] :: {it['title'] or '(no title)'}")
    else:
        print("Set RESUME_PATH to a PDF/DOCX/TXT file to test link extraction.")

if __name__ == "__main__":
    example_run()